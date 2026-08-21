# -*- coding: utf-8 -*-
"""Shared, professional DOCX styling used by BOTH the Research Report and the
Technical Documentation, so the two mentor-facing deliverables read as one
consistent, deliberately-designed pair -- not two independently-styled
scripts. No AI-generated decorative art anywhere: every image either of
these documents embeds is either a real matplotlib chart driven by
reports/*.csv, or a hand-drawn schematic box diagram of the real pipeline.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Palette (one accent family, used consistently in both documents) ----
NAVY = RGBColor(0x16, 0x2C, 0x52)      # headings, header-row fill
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)    # H2, links, accent rules
GREY = RGBColor(0x5B, 0x63, 0x70)      # captions, secondary text
INK = RGBColor(0x22, 0x25, 0x2A)       # body text
BAND = "EEF2F7"                        # banded-row table fill (hex, no #)
RULE = RGBColor(0xC9, 0xD2, 0xDE)      # thin rule / border color

BODY_FONT = "Cambria"       # serif body -- reads as a formal, printed report
HEAD_FONT = "Calibri"       # sans headings -- clean contrast against the serif body
MONO_FONT = "Consolas"


def _set_cell_border(cell, **kwargs):
    """kwargs: top/bottom/left/right/insideH/insideV -> dict(sz, color, val)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 4)))
        el.set(qn("w:color"), spec.get("color", "C9D2DE"))
        borders.append(el)
    tcPr.append(borders)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def setup_document(doc: Document, footer_label: str):
    """Base typography, margins, and a page-number + label footer."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.18
    pf.space_after = Pt(7)

    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.3)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        footer = s.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run(f"{footer_label}   |   ")
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
        r.font.name = HEAD_FONT
        _add_page_number_field(fp)
        for run in fp.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = GREY
            run.font.name = HEAD_FONT


def add_rule(doc, color="C9D2DE", size=10):
    """A thin horizontal rule via a bottom paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    h.paragraph_format.space_before = Pt(20 if level == 1 else 13)
    h.paragraph_format.space_after = Pt(9 if level == 1 else 5)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    r.font.name = HEAD_FONT
    r.font.bold = True
    r.font.size = Pt(16 if level == 1 else (12.5 if level == 2 else 11))
    r.font.color.rgb = NAVY if level == 1 else (ACCENT if level == 2 else GREY)
    if level == 1:
        add_rule(doc, color="2B6CB0", size=14)
    return h


def para(doc, text, size=10.5, italic=False, bold=False, color=None, space_after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = BODY_FONT
    r.italic = italic
    r.bold = bold
    r.font.color.rgb = color or INK
    return p


def callout(doc, text, accent=None):
    """A visually distinct one-line takeaway box (light band + left accent)."""
    accent = accent or ACCENT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F3F6FA"); pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6"); left.set(qn("w:color"), "2B6CB0")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.font.color.rgb = INK


def code_block(doc, text, size=9.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F5F6F8"); pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4"); el.set(qn("w:color"), "D7DCE3")
        pBdr.append(el)
    pPr.append(pBdr)
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line if line.strip() else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x2A, 0x2E, 0x35)
    return p


def add_table(doc, headers, rows, first_col_left=True):
    """Banded-row table, thin rules, navy header -- deliberately lighter
    than a full black grid so dense tables stay readable."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = HEAD_FONT
        r.font.bold = True
        r.font.size = Pt(9.3)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "16294F")
        _set_cell_border(hdr[i], bottom=dict(sz=6, color="16294F"))

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2 == 1:
            for c in cells:
                shade_cell(c, BAND)
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = BODY_FONT
            r.font.size = Pt(9.3)
            r.font.color.rgb = INK
            _set_cell_border(cells[i], bottom=dict(sz=2, color="E3E7EE"))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    return table


def add_figure(doc, path, caption, width=5.9, number=None):
    if not path.exists():
        para(doc, f"[Missing figure: {path.name}]", color=RGBColor(0xB0, 0x2A, 0x2A))
        return
    doc.add_picture(str(path), width=Inches(width))
    pic_p = doc.paragraphs[-1]
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    label = f"Hinh {number}. " if number else ""
    r = cap.add_run(label)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = HEAD_FONT
    r.font.color.rgb = ACCENT
    r2 = cap.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.name = BODY_FONT
    r2.font.color.rgb = GREY


def title_page(doc, kicker, title, subtitle, meta_lines):
    """A deliberately designed cover: kicker label, large title, accent
    rule, subtitle, and a small meta block -- not just a centered H1."""
    doc.add_paragraph().paragraph_format.space_after = Pt(60)

    k = doc.add_paragraph()
    k.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = k.add_run(kicker.upper())
    r.font.name = HEAD_FONT
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    k.paragraph_format.space_after = Pt(10)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.font.name = HEAD_FONT
    r.font.size = Pt(27)
    r.font.bold = True
    r.font.color.rgb = NAVY
    t.paragraph_format.space_after = Pt(14)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "2B6CB0")
    pBdr.append(bottom); rp.append(pBdr)
    rule.paragraph_format.space_after = Pt(18)
    ind = rule.paragraph_format
    ind.left_indent = Inches(1.7)
    ind.right_indent = Inches(1.7)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle)
    r.font.name = BODY_FONT
    r.italic = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = GREY
    s.paragraph_format.space_after = Pt(50)

    for line in meta_lines:
        m = doc.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = m.add_run(line)
        r.font.name = HEAD_FONT
        r.font.size = Pt(10)
        r.font.color.rgb = GREY
        m.paragraph_format.space_after = Pt(3)

    doc.add_page_break()
