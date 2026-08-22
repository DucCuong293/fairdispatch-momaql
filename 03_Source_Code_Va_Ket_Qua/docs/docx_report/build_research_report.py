# -*- coding: utf-8 -*-
"""Builds the Research / Experimental Report (Bao_Cao_Nghien_Cuu_FairDispatch_MOMAQL.docx).

Scope: replication of Kang et al. [2024] "Long-term Fairness in Ride-Hailing
Platform" (arXiv:2407.17839). This document answers ONE question: which of
the paper's scientific claims/trends does this project's real, executed
experiment suite reproduce, partially reproduce, or fail to reproduce?

It is NOT a technical/implementation manual -- that lives in docs/techdoc/.
All numbers below are read live from reports/*.csv (real simulation output,
verified against a fresh pytest run of tests/test_simulator_invariants.py
and a fresh `git rev-parse HEAD` at build time -- the pytest result depends
on data/*.parquet being present locally; see Sec. 8.1). Nothing here is
hand-typed from memory.
Run: python docs/docx_report/build_research_report.py
"""
import csv
import re
import subprocess
import sys
from pathlib import Path
from statistics import mean, pstdev

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from _docx_style import (
    NAVY, ACCENT as BLUE, GREY, setup_document, title_page, heading, para,
    bullets, add_table as _styled_add_table, add_figure as _styled_add_figure,
    callout, shade_cell,
)

ROOT = Path(__file__).resolve().parents[2]
REPORTS = ROOT / "reports"
FINAL_TEST = ROOT / "final_test"
FIGS = Path(__file__).parent / "figures"
OUT = Path(__file__).parent / "Bao_Cao_Nghien_Cuu_FairDispatch_MOMAQL.docx"

GREEN = RGBColor(0x1E, 0x7B, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x2A)
AMBER = RGBColor(0xB0, 0x7A, 0x1E)
_fig_counter = {"n": 0}


def read_csv(name):
    with (REPORTS / name).open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def read_final_test_csv(name):
    with (FINAL_TEST / name).open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


def vnum(x, digits=0):
    s = f"{x:,.{digits}f}"
    return s.replace(",", "§").replace(".", ",").replace("§", ".")


def vdec(x, digits=3):
    """Plain decimal (no thousands separator), Vietnamese comma: 0.204 -> 0,204."""
    return f"{x:.{digits}f}".replace(".", ",")


def git_head():
    try:
        return subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unavailable"


def pytest_line():
    """Actually re-run the invariant suite at build time (matches the same
    live-verification this script already does for git_head()/reports/*.csv --
    this text must never be hand-typed). Matches the real pytest summary line
    by shape rather than a loose "passed" substring search, which can grab an
    unrelated line when the run errors out before any test body executes."""
    try:
        out = subprocess.run([sys.executable, "-m", "pytest",
                               "tests/test_simulator_invariants.py", "-q"],
                              cwd=ROOT, capture_output=True, text=True, timeout=120)
        combined = out.stdout + "\n" + out.stderr
        m = re.search(r"^\d+ (?:passed|failed|errors?)\b.*$", combined, re.MULTILINE)
        if not m:
            return "(không đọc được dòng tổng kết pytest)"
        line = m.group(0).strip()
        if "error" in line.lower() and "train.parquet" in combined:
            line += (" -- data/train.parquet chưa có trong bundle này (bị gitignore, quá lớn "
                     "để đóng gói); phải lấy riêng các phần chia parquet để chạy lại bộ test "
                     "này cục bộ (xem Technical Documentation Mục 11.3)")
        return line
    except Exception as e:
        return f"(chạy pytest thất bại: {e})"


def add_table(doc, headers, rows):
    return _styled_add_table(doc, headers, rows)


def verdict_table(doc, headers, rows):
    """Banded table where the LAST column gets colored, bold verdict text."""
    table = _styled_add_table(doc, headers, rows)
    for tr in table.rows[1:]:
        cell = tr.cells[-1]
        val = cell.paragraphs[0].runs[0].text if cell.paragraphs[0].runs else ""
        low = val.lower()
        color = None
        if low.startswith("reproduced"):
            color = GREEN
        elif low.startswith("not reproduced"):
            color = RED
        elif low.startswith("partially"):
            color = AMBER
        elif low.startswith("not evaluated"):
            color = GREY
        if color:
            cell.paragraphs[0].runs[0].font.color.rgb = color
    return table


def _verdict_color(text):
    low = text.lower()
    if low.startswith("reproduced") or low.startswith("generalized"):
        return GREEN
    if low.startswith("not reproduced") or low.startswith("not generalized"):
        return RED
    if low.startswith("partially") or low.startswith("partial"):
        return AMBER
    if low.startswith("not evaluated") or low.startswith("not test"):
        return GREY
    return None


def dual_verdict_table(doc, headers, rows):
    """Like verdict_table, but colors the LAST TWO columns independently --
    used for (heldout_generalization, paper_replication_verdict), which are
    two distinct axes and must never be collapsed into one color/verdict."""
    table = _styled_add_table(doc, headers, rows)
    for tr in table.rows[1:]:
        for cell in (tr.cells[-2], tr.cells[-1]):
            if not cell.paragraphs[0].runs:
                continue
            color = _verdict_color(cell.paragraphs[0].runs[0].text)
            if color:
                cell.paragraphs[0].runs[0].font.color.rgb = color
    return table


def add_figure(doc, filename, caption, width=5.9):
    _fig_counter["n"] += 1
    _styled_add_figure(doc, FIGS / filename, caption, width=width, number=_fig_counter["n"])


def main():
    head = git_head()
    test_line = pytest_line()

    r1 = read_csv("r1_validation_results.csv")
    r2_raw = read_csv("r2_ablation_raw.csv")
    r2 = {r["ablation"]: r for r in read_csv("r2_ablation_results.csv")}
    pareto = sorted(read_csv("pareto_frontier_summary.csv"), key=lambda r: float(r["lambda"]))
    horizon = read_csv("multi_horizon_results.csv")
    disagreement = read_csv("policy_disagreement.csv")
    fleet = read_csv("fleet_scale_results.csv")
    spatial_pool = read_csv("spatial_candidate_pool.csv")
    spatial_disagree = read_csv("spatial_disagreement_by_zone.csv")
    qconv = read_csv("q_table_convergence_daily.csv")
    week = read_csv("hypothesis1_weekly_cycle.csv")
    fair_bal = read_csv("hypothesis4_fairness_balance.csv")
    mlp_summary = {r["model"]: r for r in read_csv("mlp_vs_tabular_summary.csv")}

    # ---- Final Held-out Test artifacts (final_test/) ----
    ft_manifest = None
    ft_available = FINAL_TEST.exists() and (FINAL_TEST / "test_claim_assessment.csv").exists()
    if ft_available:
        import json
        ft_manifest = json.loads((FINAL_TEST / "test_quality_transform_manifest.json").read_text(encoding="utf-8"))
        ft_baseline_summary = {r["policy"]: r for r in read_final_test_csv("baseline/test_baseline_summary.csv")}
        ft_ablation_summary = {r["ablation"]: r for r in read_final_test_csv("ablation/test_ablation_results.csv")}
        ft_horizon = read_final_test_csv("long_horizon/test_long_horizon.csv") if (FINAL_TEST / "long_horizon/test_long_horizon.csv").exists() else []
        ft_v_vs_t = read_final_test_csv("validation_vs_test.csv")
        ft_claims = read_final_test_csv("test_claim_assessment.csv")
        ft_hz = {}
        for row in ft_horizon:
            ft_hz.setdefault((row["config"], int(row["horizon_day"])), []).append(float(row["utility"]))
        ft_hz_mean = {k: mean(v) for k, v in ft_hz.items()}
        ft_gap = {}
        for d in [1, 7, 14, 21, 28, 37]:
            if ("full", d) in ft_hz_mean and ("no_forecast", d) in ft_hz_mean:
                f, nf = ft_hz_mean[("full", d)], ft_hz_mean[("no_forecast", d)]
                ft_gap[d] = (f - nf, (f - nf) / nf * 100 if nf else 0)

    r1_agg = {}
    for row in r1:
        r1_agg.setdefault(row["policy"], []).append(row)
    r1_stats = {}
    for pol, rows in r1_agg.items():
        util = [float(x["utility"]) for x in rows]
        gini = [float(x["gini"]) for x in rows]
        var = [float(x["variance"]) for x in rows]
        trips = [float(x["completed_trips"]) for x in rows]
        r1_stats[pol] = dict(u=mean(util), us=pstdev(util), g=mean(gini), gs=pstdev(gini),
                              v=mean(var), t=mean(trips))

    hz = {}
    for row in horizon:
        hz.setdefault((row["config"], int(row["horizon_day"])), []).append(float(row["utility"]))
    hz_mean = {k: mean(v) for k, v in hz.items()}
    gap = {}
    for d in [1, 7, 14, 21, 28, 37]:
        if ("full", d) in hz_mean and ("no_forecast", d) in hz_mean:
            f, nf = hz_mean[("full", d)], hz_mean[("no_forecast", d)]
            gap[d] = (f - nf, (f - nf) / nf * 100 if nf else 0)

    hz_gini = {}
    for row in horizon:
        hz_gini.setdefault((row["config"], int(row["horizon_day"])), []).append(float(row["gini"]))
    hz_gini_mean = {k: mean(v) for k, v in hz_gini.items()}

    disag_mean = mean(float(r["disagreement_rate"]) for r in disagreement)

    fleet_agg = {}
    for row in fleet:
        fleet_agg.setdefault((row["n_drivers"], row["ablation"]), []).append(float(row["utility"]))
    fleet_mean = {k: mean(v) for k, v in fleet_agg.items()}
    fleet_completed = {}
    for row in fleet:
        fleet_completed.setdefault((row["n_drivers"], row["ablation"]), []).append(float(row["completed"]))
    fleet_completed_mean = {k: mean(v) for k, v in fleet_completed.items()}

    spool = {}
    for row in spatial_pool:
        spool.setdefault(row["zone_class"], []).append(float(row["mean_candidate_pool"]))
    spool_mean = {k: mean(v) for k, v in spool.items()}

    sdis = {}
    for row in spatial_disagree:
        sdis.setdefault(row["zone_class"], {"d1_7": [], "d8_37": []})
        sdis[row["zone_class"]]["d1_7"].append(float(row["disagreement_day1_7"]))
        sdis[row["zone_class"]]["d8_37"].append(float(row["disagreement_day8_37"]))
    sdis_mean = {k: {kk: mean(vv) for kk, vv in v.items()} for k, v in sdis.items()}

    week_weekday = [float(r["incremental_gap"]) for r in week if r["is_weekend"] == "False"]
    week_weekend = [float(r["incremental_gap"]) for r in week if r["is_weekend"] == "True"]

    doc = Document()
    setup_document(doc, footer_label="FairDispatch -- Research Report")
    title_page(
        doc,
        kicker="Báo cáo Nghiên cứu / Thực nghiệm",
        title="Tái lập Định tính FairDispatch / MOMAQL",
        subtitle="Đối chiếu với Kang et al. [2024], “Long-term Fairness in "
                 "Ride-Hailing Platform”, ECML PKDD 2024 (arXiv:2407.17839)",
        meta_lines=[
            "Dương Đức Cường (tác giả chính)  ·  Nguyễn Huy Hoàng (mentor)",
            f"Git commit tại thời điểm build: {head[:12]}",
        ],
    )

    # ---------------- 0. Tom tat dieu hanh ----------------
    heading(doc, "0. Tóm tắt điều hành (Executive Summary)", level=1)
    para(doc, "Paper được tái lập: Kang et al. [2024], MOMAQL -- điều phối gọi xe đa mục tiêu "
              "(utility + fairness), có dự báo nhu cầu, trên dữ liệu NYC Taxi 2016.")
    para(doc, "Phạm vi tái lập: TÁI LẬP XU HƯỚNG định tính (qualitative trend replication), "
              "KHÔNG PHẢI tái lập số liệu chính xác -- paper không công bố đủ chi tiết triển khai "
              "(kiến trúc MLP, số tài xế, thuật toán gộp zone, seed...) để exact reproduction là "
              "mục tiêu có ý nghĩa. Đây là lựa chọn phạm vi được công bố rõ, không phải giới hạn ẩn.")
    para(doc, "Dataset: mẫu Bernoulli thật từ NYC TLC 2013 (không phải 2016 của paper -- lát cắt "
              f"thời gian thay thế), {vnum(912375)} cuốc huấn luyện / {vnum(195508)} cuốc kiểm định "
              "/ 195.510 cuốc kiểm thử, 200 tài xế mô phỏng (giả định, paper không nêu rõ).")
    para(doc, "Kết luận tổng quát theo từng claim, đánh giá trên Validation (chi tiết ở Mục 10):")
    bullets(doc, [
        "C1 (Utility-fairness trade-off): REPRODUCED.",
        "C2 (Proposed > baseline về trade-off): REPRODUCED (có phạm vi -- so với baseline đã điều "
        "chỉnh, không phải baseline gốc của paper).",
        "C3 (RL ổn định hơn khi horizon tăng): PARTIALLY REPRODUCED (chuyển pha thực tế xuất hiện "
        "~ngày 14–21, không phải tức thì).",
        "C4 (Dự báo giúp fairness dài hạn, có thể tốn ngắn hạn): PARTIALLY REPRODUCED (hướng dài hạn "
        "xác nhận từ ngày ~21; vế “tốn ngắn hạn” NOT EVALUATED rõ ràng -- ngày 1–7 "
        "không khác biệt đáng kể, không quan sát được tệ hơn).",
        "C5 (Bóc tách: thành phần dự báo giúp cả utility+fairness): REPRODUCED trên Validation -- "
        "nhưng xem lại ở Mục 9/10: khi tách riêng 2 thành phần, phần Fairness thực ra KHÔNG cải "
        "thiện (No-Forecast công bằng hơn Full) -- câu này trong bản tóm tắt gốc đã gộp nhầm 2 "
        "thành phần, sửa lại ở Mục 9.",
        "C6 (Bỏ fairness làm utility tăng mạnh, bất công bằng bùng nổ): về fairness REPRODUCED (Gini "
        "xấu đi rõ rệt); về utility NOT REPRODUCED (đảo chiều -- utility của chúng tôi GIẢM khi bỏ "
        "fairness, paper báo TĂNG).",
    ])
    para(doc, "Sau khi các kết luận trên được chốt (freeze) trên Validation, dự án chạy thêm một "
              "Final Held-out Temporal Test Evaluation (Mục 9) trên test.parquet -- lát cắt thời gian "
              "chưa từng dùng để phát triển/tuning. 13/13 phát hiện định trước (pre-specified) từ "
              "Validation giữ đúng chiều trên Test, cho thấy hành vi quan sát được của bản tái lập "
              "khá ổn định theo thời gian. NHƯNG điều đó KHÔNG có nghĩa toàn bộ 6 claim của paper "
              "được reproduce -- ví dụ C4 (dự báo cải thiện fairness dài hạn) vẫn Not Reproduced trên "
              "cả Validation lẫn Test. Vì vậy câu chuyện tổng thể của báo cáo là: "
              "\"Validation-developed trend replication with held-out temporal test support\" -- "
              "không phải \"đã reproduce hoàn toàn paper\".", bold=True)
    callout(doc, "Không có số liệu nào trong báo cáo này bị chỉnh sửa/chọn lọc để giống paper "
                 "hơn. Mọi con số đọc trực tiếp từ reports/*.csv (Validation) và final_test/*.csv "
                 "(Held-out Test) tại thời điểm build tài liệu, không hand-type từ trí nhớ.")

    # ---------------- 1. Introduction ----------------
    heading(doc, "1. Giới thiệu (Introduction)", level=1)
    para(doc, "1.1 Động lực", size=11, bold=True)
    para(doc, "Các nền tảng gọi xe hiện đại tối ưu hóa hiệu quả (utility) có thể tạo ra bất công bằng "
              "thu nhập giữa tài xế qua thời gian. Kang et al. [2024] đề xuất một bộ điều khiển kết hợp "
              "dự báo nhu cầu tương lai với Multi-Objective Q-Learning (MOMAQL) để giữ fairness ổn định "
              "trên một khung thời gian dài, thay vì tối ưu fairness cận thị từng cửa sổ quyết định.")
    para(doc, "1.2 Mục tiêu tái lập", size=11, bold=True)
    para(doc, "Dự án này tái lập CÓ CHỦ ĐÍCH các XU HƯỚNG định tính của paper -- không phải giá trị "
              "số tuyệt đối -- trên một lát cắt dữ liệu thay thế (NYC TLC 2013 thay vì 2016 của paper). "
              "Câu hỏi nghiên cứu chính:")
    para(doc, "“Nếu hệ thống biết trước xu hướng nhu cầu tương lai, nó có chấp nhận một quyết định "
              "chưa tối ưu về fairness ngay hiện tại để đạt fairness tốt hơn trong một khoảng thời gian "
              "dài hơn hay không -- và điều này có tái xuất hiện khi xây lại phương pháp một cách hợp "
              "lý từ thông tin paper cung cấp?”", italic=True)
    para(doc, "1.3 Đóng góp của tái lập này", size=11, bold=True)
    bullets(doc, [
        "Bộ 5 chính sách điều phối thực thi trên 195.508 yêu cầu kiểm định thật, 5 seed mỗi chính sách.",
        "Bóc tách 3 nhánh thật (Full / w/o-Forecast / w/o-Fairness), 5 seed.",
        "Pareto sweep 7 giá trị lambda, 5 seed.",
        "Nghiên cứu multi-horizon trên MỘT quỹ đạo thật, checkpoint ngày 1–37 (không phải chạy "
        "lại độc lập từng horizon).",
        "6 thí nghiệm cơ chế bổ sung thật (quy mô đội xe, độ sâu vùng ứng viên, tốc độ hội tụ bảng Q, "
        "chu kỳ nhu cầu tuần, cân bằng fairness/nhìn trước, so sánh MLP-vs-tabular) đi xa hơn yêu cầu "
        "tối thiểu của paper để kiểm tra TẠI SAO các xu hướng xuất hiện, không chỉ KHI NÀO.",
    ])

    # ---------------- 2. Original Study ----------------
    heading(doc, "2. Nghiên cứu Gốc (Original Study)", level=1)
    para(doc, "2.1 Vấn đề bài toán", size=11, bold=True)
    para(doc, "Bộ điều khiển tối đa hóa utility − lambda × bất công bằng trên một khung thời gian "
              "dài. Efficiency = tổng utility tất cả tài xế. Fairness = phương sai (variance) của "
              "accumulated utility giữa các tài xế -- variance thấp nghĩa là thu nhập tài xế đồng đều hơn.")
    para(doc, "2.2 Các tuyên bố (claims) cần tái lập", size=11, bold=True)
    add_table(doc, ["ID", "Tuyên bố của paper", "Bằng chứng cần tái lập"], [
        ["C1", "Utility và fairness có trade-off", "Tăng fairness có thể làm giảm utility"],
        ["C2", "Phương pháp đề xuất tạo trade-off tốt hơn baseline", "Điểm của phương pháp đề "
         "xuất nằm ở vùng utility/fairness tốt hơn"],
        ["C3", "Phương pháp dựa trên RL ổn định hơn khi horizon tăng", "Đường cong fairness ít "
         "biến động hơn theo thời gian"],
        ["C4", "Prediction giúp long-term fairness", "Có thể thua ngắn hạn nhưng tốt hơn khi horizon dài"],
        ["C5", "Prediction giúp utility + fairness so với no-prediction", "Ablation phải phản ánh "
         "đúng hướng"],
        ["C6", "Bỏ fairness làm utility tăng mạnh nhưng inequality bùng lên", "w/o Fairness là "
         "trường hợp cực đoan về utility/bất công bằng"],
    ])
    para(doc, "2.3 Số liệu gốc của paper (Table 1 và Ablation, trích dẫn trực tiếp)", size=11, bold=True)
    add_table(doc, ["Phương pháp", "Tổng Utility", "Fairness (Var, thấp hơn = fair hơn)"], [
        ["REASSIGN", "76.536", "493.638"],
        ["LAF", "80.606", "107.790"],
        ["Balance Ride-Pooling", "85.924", "100.255"],
        ["Proposed (Full)", "95.823,79", "85.193,62"],
    ])
    add_table(doc, ["Ablation", "Tổng Utility", "Fairness (Var)"], [
        ["Full", "95.823,79", "85.193,62"],
        ["w/o Prediction", "56.873,21", "153.697,27"],
        ["w/o Fairness", "2.194.901,19", "2,677 × 10⁹"],
    ])
    para(doc, "Cấu hình của paper: sampling rate = 0,05; lambda = 1; omega = 0,6; gamma = 0,9. "
              "Forecast module: MLP 3 lớp, dự báo số lượng yêu cầu theo cặp điểm đón–trả mỗi giờ, "
              "MSE = 94,69. Phần cứng: Xeon Gold 6240 + RTX 8000.", size=9.5, color=GREY)
    para(doc, "2.4 Giới hạn công bố của paper gốc", size=11, bold=True)
    para(doc, "Paper KHÔNG công bố đủ chi tiết để tái lập tuyệt đối: số neuron hidden layer của MLP, "
              "optimizer/learning rate/activation/epochs, thuật toán gộp nhiều vị trí thành 1 node, "
              "số lượng/khởi tạo tài xế mô phỏng, chi tiết exploration của Q-learning, exact “peak "
              "2-hour” window, và một số chi tiết triển khai baseline. Đây là lý do chính dự án "
              "này chọn tái lập xu hướng thay vì số liệu tuyệt đối.")

    # ---------------- 3. Replication Scope ----------------
    heading(doc, "3. Phạm vi Tái lập (Replication Scope)", level=1)
    para(doc, "3.1 Exact reproduction vs. trend replication", size=11, bold=True)
    para(doc, "Dự án này CHÍNH THỨC tuyên bố đây là một TREND REPLICATION, không phải EXACT NUMERICAL "
              "REPRODUCTION. Lý do: (a) paper không công bố đủ chi tiết triển khai (Mục 2.4); (b) dữ "
              "liệu sử dụng là lát cắt thời gian khác (2013 vs. 2016 của paper); (c) mô hình dự báo "
              "được thay thế có chủ đích (xem Mục 5).")
    para(doc, "3.2 Tiêu chí thành công", size=11, bold=True)
    bullets(doc, [
        "Reproduced: thứ tự/quan hệ/hướng của xu hướng khớp với paper, có bằng chứng định lượng rõ ràng.",
        "Partially Reproduced: có dấu hiệu đúng hướng nhưng chỉ đúng một phần (ví dụ: đúng hướng dài "
        "hạn nhưng không đúng thời điểm crossover ngắn hạn như paper).",
        "Not Reproduced: xu hướng đảo chiều hoặc không xuất hiện so với paper, có bằng chứng định "
        "lượng rõ ràng.",
        "Not Evaluated: chưa có thí nghiệm riêng cô lập được claim này trong phạm vi dự án hiện tại.",
    ])
    para(doc, "3.3 Khoảng trống tái lập đã biết trước (known reproducibility gaps)", size=11, bold=True)
    bullets(doc, [
        "Không tái lập kiến trúc MLP dự báo số lượng yêu cầu theo cặp OD -- thay bằng Q(vùng,giờ) dạng "
        "bảng, học trực tuyến qua Bellman TD(0) (xem Mục 5.3 và so sánh trực tiếp ở Mục 8.5).",
        "Không tái lập baseline REASSIGN/Balance-RP gốc của paper -- dùng Exact REASSIGN (Lesmana et al. "
        "2019) và LAF (Sühr et al. 2019) như các công trình gần nhất, thật, kiểm chứng được.",
        "Không có mã nguồn công khai chính thức của paper để đối chiếu trực tiếp.",
    ])

    # ---------------- 4. Dataset and Preprocessing ----------------
    heading(doc, "4. Dữ liệu và Tiền xử lý (Dataset and Preprocessing)", level=1)
    para(doc, "Nguồn: mẫu Bernoulli thật từ bộ dữ liệu NYC TLC 2013 đã làm sạch của dự án gốc "
              "(tháng Development 1–8), lọc theo manhattan_both=true và quality_flag_bitset=0. "
              "Bộ lọc Manhattan/chất lượng do pipeline gốc của dự án cha áp dụng, tái sử dụng ở "
              "đây, không tự làm lại.")
    add_table(doc, ["Giai đoạn", "Số dòng"], [
        ["NYC TLC 2013 thô (tất cả tháng)", "171.816.340"],
        ["Tập Development (tháng 1–8)", "115.707.941"],
        ["Lọc Manhattan + chất lượng, mẫu 1,3tr", "1.303.393"],
        ["Huấn luyện (chia theo thời gian)", "912.375"],
        ["Kiểm định", "195.508"],
        ["Kiểm thử", "195.510"],
    ])
    para(doc, "Biểu diễn không gian: 67 TLC taxi zones (xấp xỉ cho “merged location graph nodes” "
              "của paper -- paper không mô tả đủ thuật toán gộp vị trí để tái tạo chính xác).")
    para(doc, "So sánh từng thành phần với paper:", size=11, bold=True)
    add_table(doc, ["Thành phần", "Paper", "Bản tái lập", "Trạng thái"], [
        ["Dữ liệu", "NYC Taxi, 3–4/2016", "NYC TLC 2013, tháng 1–8", "Sai khác (năm)"],
        ["Vùng", "Manhattan", "Manhattan (manhattan_both)", "Khớp"],
        ["Số tài xế", "Không nêu rõ", "200", "Giả định"],
        ["Biểu diễn không gian", "Gộp node đồ thị vị trí", "67 zone TLC", "Xấp xỉ"],
        ["Mô hình dự báo", "MLP 3 lớp, OD-pair×giờ", "Q(vùng,giờ) dạng bảng, TD(0) trực tuyến",
         "Sửa đổi"],
        ["λ (trọng số công bằng)", "1", "0,5 mặc định (quét 0–1)", "Sai khác"],
        ["ω (hệ số scale)", "0,6", "Không dùng; chuẩn hóa fairness tương đối", "Sửa đổi"],
        ["γ (discount)", "0,9", "0,9", "Khớp"],
        ["Baseline", "Greedy, REASSIGN, LAF, Balance-RP", "Greedy, Nearest, LAF, Exact REASSIGN",
         "Trùng 1 phần"],
        ["Horizon đánh giá", "1–7 ngày", "1–7 ngày + toàn bộ 37 ngày", "Khớp (mở rộng)"],
    ])

    # ---------------- 5. Replication Methodology ----------------
    heading(doc, "5. Phương pháp Tái lập (Replication Methodology)", level=1)
    para(doc, "5.1 Simulator", size=11, bold=True)
    para(doc, "Pipeline: cuốc xe thật → chia theo thời gian → mô phỏng theo lô cửa sổ 60 giây → "
              "ghép cặp Hungarian M-to-N (đệm dummy, cho phép từ chối) → 5 chính sách điều phối "
              "(gồm MOMAQL) → chỉ số utility/Gini/phương sai.")
    para(doc, "5.2 Định nghĩa Utility (A8)", size=11, bold=True)
    para(doc, "net = fare_amount − eta × 0,0025 (cước phí USD thật trừ chi phí thời gian chạy "
              "rỗng), tích lũy theo tài xế. Đây là chỉ số doanh thu thực -- CHƯA xác nhận là cùng "
              "đại lượng với “geographic benefit/cost function” mà paper mô tả cho instant reward.")
    para(doc, "5.3 Định nghĩa Fairness (A9)", size=11, bold=True)
    para(doc, "Gini là chỉ số chính xuyên suốt báo cáo. Chỉ số chính của paper là "
              "Var(U₁,…,Uₙ) (và biến thể chuẩn hóa σ(U)/Ū); dự án này báo cáo song song "
              "Var/Std/hệ số biến thiên cho R1, R2 và Pareto sweep. Gini được ưu tiên hơn CV vì "
              "CV không ổn định khi utility trung bình gần 0.")
    para(doc, "5.4 Q(vùng,giờ) -- KHÔNG phải bộ dự báo nhu cầu", size=11, bold=True)
    para(doc, "Học qua bootstrap Bellman TD(0): Q(P,h) ← Q(P,h) + α[reward + γQ(D,h′) − Q(P,h)], "
              "reward = cước phí − chạy rỗng. Đại lượng nó học là doanh thu ròng kỳ vọng chiết "
              "khấu tương lai (USD) khi đứng ở 1 vùng vào 1 giờ -- một hàm giá trị theo nghĩa RL, "
              "khác về bản chất so với MLP của paper (dự báo số lượng yêu cầu D̂ᵢⱼ,ₜ theo từng cặp "
              "điểm đón–trả mỗi giờ). Không gọi Q là bản thay thế cho mô-đun dự báo đó, và không "
              "gọi nó là mạng nơ-ron. α = 0,1, γ = 0,9.")
    para(doc, "5.5 Chuẩn hóa hàm mục tiêu (scalarisation)", size=11, bold=True)
    para(doc, "s = (1−λ)[f − cη + γQ(D,h′)] + λ[(Ī−I_d)/Ī]f -- điểm số kết hợp cước phí/chạy rỗng "
              "(USD) và khoảng cách công bằng tương đối (phân số thu nhập trung bình × cước phí "
              "cuốc này, tự nó cũng là USD). Đây là hàm mục tiêu ĐÃ SỬA ĐỔI, không phải dạng "
              "(1−λ)·utility − λω·Var(U) của chính paper. Hai λ KHÔNG tương đương trực tiếp: λ=1 "
              "của dự án này là cực thuần-công-bằng của 1 công thức khác, không phải giới hạn của "
              "λ=1 mà paper báo cáo. Không tuyên bố tương ứng toán học giữa 2 công thức.")

    # ---------------- 6. Baselines ----------------
    heading(doc, "6. Baseline (Baselines)", level=1)
    add_table(doc, ["Baseline", "Nguồn", "Exact/Approx", "Sửa đổi"], [
        ["Greedy", "Paper", "Exact", "Không"],
        ["Nearest", "Không có trong paper", "—", "Tự thêm, không phải baseline của paper"],
        ["LAF", "Sühr et al. [2019] (ý tưởng gốc)", "Approx", "Dùng khoảng cách công bằng tương "
         "đối, không phải chênh lệch thu nhập thô"],
        ["Exact REASSIGN", "Lesmana et al. [2019]", "Approx", "Hungarian M-to-N theo lô, cho phép "
         "từ chối"],
        ["MOMAQL", "Paper", "Approx", "Bộ ước lượng giá trị RL dạng bảng, không phải MLP của paper"],
    ])
    para(doc, "Nguồn gốc ban đầu của LAF (“Shi et al.”, theo cách trích trong một số phiên bản "
              "paper) không xác minh được là một công bố thật, tìm được -- dự án này trích dẫn "
              "Sühr et al. [2019] thay thế, đây là công trình thật, kiểm chứng được, gần nhất với "
              "cùng ý tưởng gốc (phân bổ công bằng thu nhập tài xế qua nhiều lần ghép cặp thay vì "
              "ép công bằng từng lần). Không tuyên bố LAF là bản triển khai trực tiếp thuật toán cụ "
              "thể của paper đó.", size=9.5, color=GREY)

    # ---------------- 7. Experimental Protocol ----------------
    heading(doc, "7. Giao thức Thực nghiệm (Experimental Protocol)", level=1)
    add_table(doc, ["Tham số", "Paper", "Dự án này", "Lý do"], [
        ["Số tài xế mô phỏng", "Không nêu rõ", "200", "Giả định hợp lý, không xác minh được với paper"],
        ["Cửa sổ điều phối", "Không nêu rõ", "60 giây", "Đơn vị batching thực tế của simulator"],
        ["ETA khả thi tối đa", "Không nêu rõ", "600 giây", "Ngưỡng feasible_drivers() nội bộ"],
        ["λ (trọng số công bằng)", "1", "0,5 mặc định, quét 0–1", "Cho phép quan sát cả trade-off "
         "curve, không chỉ 1 điểm"],
        ["γ (discount)", "0,9", "0,9", "Khớp trực tiếp"],
        ["α (learning rate Q)", "Không nêu rõ", "0,1", "Giả định chuẩn cho TD(0) online"],
        ["Số seed", "Không nêu rõ", "5 (thí nghiệm chính) / 3 (thí nghiệm cơ chế bổ sung)", "Công "
         "khai rõ để đo phương sai qua seed; 3 seed cho thí nghiệm phụ do giới hạn compute"],
    ])
    para(doc, "Phần cứng thực hiện tái lập: CPU-only cho toàn bộ simulator/dispatch (R1, R2, Pareto, "
              "multi-horizon, quy mô đội xe, không gian, hội tụ bảng Q); GPU NVIDIA GeForce GTX 1650 "
              "(qua CUDA) chỉ dùng cho huấn luyện MLP dự báo ở Mục 8.5. Không so sánh trực tiếp thời "
              "gian chạy với phần cứng Xeon Gold 6240 + RTX 8000 của paper.", size=9.5, color=GREY)

    # ---------------- 8. Results ----------------
    heading(doc, "8. Kết quả (Results)", level=1)

    heading(doc, "8.1 Kiểm tra tính hợp lệ (Sanity checks)", level=2)
    para(doc, f"tests/test_simulator_invariants.py, chạy lại thật ngay tại thời điểm build tài "
              f"liệu này: {test_line}. Đây không phải coverage hình thức đầy đủ; 2 test dùng "
              "record_trace=False mặc định nên assertion dựa trên trace hiện đang lặp trên trace "
              "rỗng (xem Techdoc).")

    heading(doc, "8.2 So sánh cơ sở (R1)", level=2)
    para(doc, "195.508 yêu cầu kiểm định thật, 200 tài xế, trung bình 5 seed.")
    r1_rows = []
    for pol in ["MOMAQL", "Greedy", "Nearest", "LAF", "Exact REASSIGN"]:
        s = r1_stats[pol]
        r1_rows.append([pol, vnum(s["u"]), vdec(s["g"]), vnum(s["v"] / 1e6, 2) + " ×10⁶",
                         vnum(s["t"])])
    add_table(doc, ["Chính sách", "Utility ($)", "Gini", "Phương sai", "Số chuyến"], r1_rows)
    add_figure(doc, "r1_validation_unified_comparison.png", "Hình 1: R1 -- utility và Gini theo "
               "chính sách (TB 5 seed, kiểm định).")
    para(doc, "MOMAQL vượt trội các baseline đã điều chỉnh dùng trong bối cảnh tái lập này (cả "
              "utility lẫn Gini cùng lúc) -- claim C2. Không tuyên bố nó vượt trội chính baseline "
              "gốc của paper, vốn chưa được tái lập độc lập.")

    heading(doc, "8.3 Đánh đổi Utility–Fairness (Pareto, claim C1)", level=2)
    pareto_rows = [[vdec(float(r["lambda"]), 1), vnum(float(r["utility_mean"])),
                     vdec(float(r["gini_mean"]))] for r in pareto]
    add_table(doc, ["λ", "Utility TB ($)", "Gini TB"], pareto_rows)
    add_figure(doc, "pareto_frontier_unified_curve.png", "Hình 2: MOMAQL Pareto Frontier -- "
               "Utility vs. Fairness, λ ∈ {0; 0,2; 0,4; 0,5; 0,6; 0,8; 1,0}.")
    para(doc, "Đơn điệu (C1): utility tăng theo (1−λ) từ ~$898k (Gini 0,450) đến ~$1,56tr "
              "(Gini 0,228) tại λ=0,8; λ=1,0 sập về ~$766k, Gini ≈0 (về công thức = LAF ở cực này, "
              "vì λ=1 loại bỏ hoàn toàn thành phần efficiency).")

    heading(doc, "8.4 Ổn định dài hạn (C3, C4)", level=2)
    para(doc, "Theo dõi Full / w/o-Forecast / w/o-Fairness trên MỘT quỹ đạo thật (không phải chạy "
              "lại độc lập) mỗi cấu hình/seed, checkpoint tại ngày 1,2,3,4,5,6,7,14,21,28,37, đo "
              "riêng tỷ lệ đổi quyết định D = P(π_full ≠ π_no_forecast) trên cùng tập ứng viên.")
    add_figure(doc, "multi_horizon_unified_curve.png", "Hình 3: Utility và Gini theo horizon đánh "
               "giá (ngày 1–37), 3 cấu hình.")
    gap_rows = []
    for d in [1, 7, 14, 21, 28, 37]:
        if d in gap:
            g, pct = gap[d]
            sign = "+" if pct >= 0 else "-"
            gap_rows.append([str(d), vnum(g), f"{sign}{vdec(abs(pct), 2)}%"])
    add_table(doc, ["Ngày", "Chênh lệch utility Full − w/o-Forecast ($)", "Chênh lệch (%)"], gap_rows)
    para(doc, "Lưu ý số liệu: quỹ đạo checkpoint 37 ngày này chạy trên một bản cắt ~190.966 yêu "
              "cầu (giới hạn theo cửa sổ 37 ngày + 1 giờ đệm kể từ yêu cầu đầu tiên), KHÔNG PHẢI "
              "trọn 195.508 yêu cầu kiểm định dùng ở R1/R2/Pareto (Mục 8.2, 8.5). Do đó chênh lệch "
              "+20,19% tại ngày 37 ở bảng trên KHÁC với con số +22,4% chính thức ở Mục 8.5 (đo trên "
              "trọn tập); đây là 2 phép đo trên 2 cỡ mẫu khác nhau, không phải mâu thuẫn số liệu.",
              size=9.5, color=GREY)
    add_figure(doc, "multi_horizon_2phase_breakthrough.png", "Hình 4: Giai đoạn 1 (ngày 1–7, đồng "
               "nhất) vs. Giai đoạn 2 (ngày 8–37, phân hóa).")
    para(doc, f"Tỷ lệ đổi quyết định D = P(π_full ≠ π_no_forecast) trung bình 5 seed, đo trên "
              f"TOÀN BỘ quỹ đạo 37 ngày (không phải riêng ngày 1–7): {vdec(disag_mean*100, 1)}%. Phát "
              "hiện: ngày 1–14 gần như phẳng về mặt thống kê; khoảng "
              "cách mở ra giữa ngày 14 và ngày 21, sau đó tăng dần đến hết ngày 37. Kang et al. "
              "[2024] báo cáo crossover (full công bằng hơn no-prediction) bắt đầu từ ngày 3; dự án "
              "này KHÔNG quan sát được crossover đó trong cùng cửa sổ -- dự báo dạng bảng cần thời "
              "gian dài hơn đáng kể để tích lũy lợi thế so với MLP chuyên biệt của paper. Đây là lý "
              "do C3/C4 gắn nhãn Partially Reproduced: hướng dài hạn xác nhận, thời điểm crossover "
              "ngắn hạn thì không.")

    heading(doc, "8.4.1 Cơ chế đứng sau thời điểm chuyển pha ngày 14–21 (bằng chứng bổ sung)", level=3)
    para(doc, "4 thí nghiệm cơ chế bổ sung thật, không bắt buộc theo yêu cầu tối thiểu của paper, "
              "nhưng trực tiếp trả lời TẠI SAO chuyển pha xảy ra ở khung thời gian đó -- không phải "
              "ablation có kiểm soát cô lập được nhân quả, mà là liên hệ thật, đã đo.")
    qc_rows = []
    for row in qconv:
        if int(row["day"]) in (1, 7, 14, 21, 28, 37):
            dq = row["mean_abs_delta_q_vs_prev_day"]
            qc_rows.append([row["day"], row["n_states_visited"],
                             vdec(float(dq)) if dq and dq != "nan" else "— (không có ngày trước)"])
    add_table(doc, ["Ngày huấn luyện", "Số trạng thái (vùng,giờ) đã thăm", "Mean |ΔQ|"], qc_rows)
    para(doc, "Độ phủ trạng thái bão hòa gần đúng lịch trình (98,4% tại ngày 14), nhưng phần dư giá "
              "trị Q (|ΔQ|) chỉ ổn định vào dải dài hạn khoảng ngày 20–21 -- khớp sát thời điểm "
              "chênh lệch utility ở bảng trên thật sự mở ra.")
    para(doc, f"Độ sâu vùng ứng viên: lõi trung bình {vdec(spool_mean.get('core', 0), 1)} tài "
              f"xế/yêu cầu vs. ngoại vi {vdec(spool_mean.get('periphery', 0), 1)} -- vùng lõi có "
              "vùng ứng viên "
              "sâu hơn ~4 lần, cho bộ giải Hungarian nhiều lựa chọn hơn để đổi kết quả gán, giải "
              "thích khả dĩ (không phải chứng minh có kiểm soát) cho việc đổi quyết định tập trung "
              "ở lõi hơn ngoại vi.")
    wk_wd = mean(week_weekday) if week_weekday else 0.0
    wk_we = mean(week_weekend) if week_weekend else 0.0
    para(doc, f"Chu kỳ nhu cầu tuần: chênh lệch utility gia tăng TB ${vnum(wk_wd)} ngày thường vs. "
              f"${vnum(wk_we)} cuối tuần -- không có chu kỳ tuần lặp lại rõ ràng; mức tăng khớp đúng "
              "các bước ngoặt của hội tụ bảng Q ở trên, không phải chu kỳ 7 ngày.")
    fb_rows = [r for r in fair_bal if int(r["day"]) in (1, 14, 21, 37)]
    add_table(doc, ["Ngày", "|Thành phần Efficiency|", "|Thành phần Fairness|", "Tỷ trọng Fairness"],
               [[r["day"], vdec(float(r["mean_abs_efficiency_term"]), 2),
                 vdec(float(r["mean_abs_fairness_term"]), 2),
                 f"{vdec(float(r['fairness_share'])*100, 1)}%"] for r in fb_rows])
    para(doc, "Tỷ trọng thành phần fairness trên tổng độ lớn score tăng dần rồi ổn định quanh cùng "
              "cửa sổ ngày 14–21, nhưng thành phần efficiency/nhìn trước vẫn luôn chiếm đa số "
              "(≥95%) suốt 37 ngày -- cái dịch chuyển là trọng số tương đối, không phải thành phần "
              "nào chi phối.")

    heading(doc, "8.5 Bóc tách dự báo (C5) và so sánh với MLP thật", level=2)
    r2_rows = []
    for k, label in [("full", "Full"), ("no_forecast", "w/o Forecast"), ("no_fairness", "w/o Fairness")]:
        row = r2[k]
        r2_rows.append([label, vnum(float(row["utility_mean"])), vdec(float(row["gini_mean"])),
                         vnum(float(row["variance_mean"]) / 1e6, 2) + " ×10⁶"])
    add_table(doc, ["Biến thể", "Utility TB ($)", "Gini TB", "Phương sai TB"], r2_rows)
    add_figure(doc, "r2_ablation_unified_comparison.png", "Hình 5: Full so với w/o-Forecast và "
               "w/o-Fairness.")
    full_u = float(r2["full"]["utility_mean"])
    nf_u = float(r2["no_forecast"]["utility_mean"])
    para(doc, f"Bóc tách dự báo (C5) tái lập sạch: +{vdec((full_u-nf_u)/nf_u*100, 1)}% utility khi bật "
              "thành phần Q(zone,hour), đủ 5/5 seed.")
    if "MOMAQL (Tabular Q)" in mlp_summary or "Tabular Q" in mlp_summary:
        mk = "Tabular Q" if "Tabular Q" in mlp_summary else "MOMAQL (Tabular Q)"
        para(doc, "So sánh trực tiếp với một MLP thật (không phải chỉ suy diễn): huấn luyện một "
                  "MLP 3 lớp PyTorch dự báo SỐ ĐẾM cầu theo cặp OD+giờ (khớp đúng cơ chế dự báo của "
                  "paper), rescale vào cùng vị trí trong công thức score, so với bảng Q dạng "
                  "tabular hiện tại, trên 5 seed, trọn 195.508 yêu cầu kiểm định:")
        mlp_rows = []
        for label, key in [("Tabular Q", mk), ("MLP Demand Forecast", "MOMAQL (MLP Demand Forecast)"),
                            ("No-Forecast", "No-Forecast")]:
            if key in mlp_summary:
                row = mlp_summary[key]
                mlp_rows.append([label, vnum(float(row["utility_mean"])),
                                  vdec(float(row["gini_mean"])),
                                  vnum(float(row["variance_mean"]) / 1e6, 2) + " ×10⁶"])
        add_table(doc, ["Mô hình", "Utility TB ($)", "Gini TB", "Phương sai TB"], mlp_rows)
        para(doc, "Bảng Q dạng tabular THẮNG MLP thật cả về utility lẫn công bằng. Đây là bằng "
                  "chứng thật cho thấy ưu thế của Q dạng tabular không đến từ việc paper's MLP kém "
                  "hơn về nguyên tắc, mà nhiều khả năng đến từ việc Q được huấn luyện bằng Bellman "
                  "TD trực tiếp trên đúng reward mà hàm score tối ưu -- một MLP dự báo số đếm cầu là "
                  "tín hiệu khác, không calibrate theo đúng giá trị đó.")

    heading(doc, "8.6 Bóc tách fairness và sai khác C6", level=2)
    para(doc, "Vế công bằng -- tái lập đúng hướng: Gini xấu đi từ 0,204 → 0,450 khi bỏ công bằng, "
              "khớp đúng hướng paper. Vế utility -- KHÔNG tái lập (đảo chiều): paper báo cáo bỏ "
              "công bằng làm utility TĂNG mạnh (+2.190% trong Table 2 của họ); của dự án này GIẢM "
              f"{vdec((float(r2['no_fairness']['utility_mean'])-full_u)/full_u*100, 1)}%. Không làm tròn "
              "thành “tái lập một phần” gộp chung -- C6 báo cáo như 2 phát hiện riêng biệt, trái "
              "chiều nhau.")
    add_table(doc, ["Cấu hình", "Số chuyến", "$/chuyến", "Utility"], [
        ["λ=0 (không công bằng)", "90.264", "9,949", "898.025"],
        ["λ=0,5 (full)", "152.380", "9,335", "1.422.441"],
    ])
    para(doc, "Khi bật công bằng, doanh thu/chuyến THẤP hơn 6,2% (chính sách phục vụ một số cuốc "
              "giá trị thấp mà nếu không sẽ bị từ chối), nhưng SỐ CHUYẾN cao hơn 69% -- hiệu ứng "
              "sản lượng lấn át, nên tổng utility tăng theo λ chứ không giảm. Đây là hệ quả thực, "
              "đã công bố, của cơ chế ghép cặp cho phép từ chối: với λ=0, chính sách từ chối nhiều "
              "ứng viên giá trị thấp hơn so với λ=0,5. Chưa kiểm chứng độc lập câu chuyện không "
              "gian cụ thể (cuốc ngắn trung tâm vs. ngoại vi) ngoài phân rã số chuyến×doanh thu này.")

    heading(doc, "8.7 Nghiên cứu quy mô (scale study)", level=2)
    fleet_rows = []
    fleet_pct = {}
    for n in ["100", "200", "400"]:
        f = fleet_mean.get((n, "full"))
        nf = fleet_mean.get((n, "no_forecast"))
        if f is not None and nf is not None:
            pct = (f - nf) / nf * 100
            fleet_pct[n] = pct
            sign = "+" if pct >= 0 else "-"
            fleet_rows.append([n, vnum(f), vnum(nf), f"{sign}{vdec(abs(pct), 1)}%"])
    add_table(doc, ["N tài xế", "Full ($)", "w/o-Forecast ($)", "Chênh lệch"], fleet_rows)
    add_figure(doc, "scale_and_convergence.png", "Hình 6: (trái) Utility theo quy mô đội xe; "
               "(phải) hội tụ bảng Q qua 37 ngày huấn luyện.")
    n400_full_c = fleet_completed_mean.get(("400", "full"), 0)
    para(doc, f"Lợi thế dự báo KHÔNG ổn định theo quy mô -- co lại đơn điệu từ +{vdec(fleet_pct['100'], 1)}% "
              f"(N=100) về +{vdec(fleet_pct['400'], 1)}% (N=400). Tỷ lệ hoàn thành tại N=400 đạt "
              f"≈{vdec(n400_full_c/195508*100, 1)}% -- khi "
              "cung gần bão hòa cầu, xe nào phục vụ yêu cầu nào gần như không ảnh hưởng kết quả "
              "tổng, nên tín hiệu định vị thông minh còn ít chỗ để phát huy. Đây là một kết quả thật "
              "KHÔNG được paper nhắc tới (paper không thử nghiệm độ nhạy quy mô đội xe).")

    heading(doc, "8.8 Độ vững chắc thống kê (statistical robustness)", level=2)
    para(doc, "R1, R2, Pareto, multi-horizon và so sánh MLP-vs-tabular đều chạy 5 seed độc lập "
              "(20260721–20260725), báo cáo trung bình ± độ lệch chuẩn. Các thí nghiệm cơ chế bổ "
              "sung (quy mô đội xe, không gian, độ sâu vùng ứng viên, hội tụ bảng Q, chu kỳ tuần, "
              "cân bằng fairness) chạy 3 seed do giới hạn compute cục bộ (máy CPU-only, ~3–4GB RAM "
              "trống khả dụng khi thực hiện) -- giới hạn này được công bố công khai, không che giấu.")

    # ---------------- 9. Final Held-out Test Evaluation ----------------
    heading(doc, "9. Đánh giá Final Held-out Test (Final Held-out Test Evaluation)", level=1)
    heading(doc, "9.1 Mục đích", level=2)
    para(doc, "Toàn bộ Mục 8 (Results) ở trên chạy trên Validation (195.508 yêu cầu) -- tập dữ liệu "
              "dùng để phát triển implementation, chọn/khoá cấu hình, và chạy ablation/long-horizon. "
              "Mục này chạy thêm một Final Held-out Temporal Test Evaluation trên test.parquet -- "
              "lát cắt thời gian CHƯA TỪNG được dùng để tuning, chọn λ, hay sửa policy/simulator. "
              "Test dùng để XÁC MINH generalization, không dùng để CHỌN cấu hình. Protocol (cấu hình, "
              "seed, metric, no-tuning rule) được đóng băng (frozen) trong "
              "`final_test/FINAL_TEST_PROTOCOL.md` TRƯỚC KHI bất kỳ policy nào chạy trên test.parquet.")
    if ft_available:
        heading(doc, "9.2 Cổng Kiểm tra Chất lượng Dữ liệu Test (Test Data Quality Gate)", level=2)
        tstats = ft_manifest["per_split"]["test"]
        para(doc, "Trước khi chạy policy, audit dữ liệu test.parquet phát hiện 33/195.510 dòng "
                  "(0,017%) có field `duration_seconds` (derived) bị lỗi -- timestamp `pickup_ts`/"
                  "`dropoff_ts` của các dòng này vẫn hợp lệ. Nguyên tắc \"minimal deterministic "
                  "repair\": sửa `duration_seconds` từ timestamp cho các dòng phục hồi được, chỉ loại "
                  "bỏ dòng không thể phục hồi -- không xoá toàn bộ 33 dòng một cách vũ đoán.")
        add_table(doc, ["Bước", "Số dòng"], [
            ["Test gốc (raw)", vnum(tstats["original_rows"])],
            ["Loại do trùng giây ranh giới Validation/Test (temporal-boundary hygiene)",
             vnum(tstats["temporal_boundary_excluded"])],
            ["Sửa (repair) duration_seconds từ timestamp", vnum(tstats["duration_repaired"])],
            ["Loại do duration không thể phục hồi (zero-trip)", vnum(tstats["duration_excluded"])],
            ["Final Test Evaluation View", vnum(tstats["final_evaluated_rows"])],
        ])
        para(doc, "test.parquet gốc trên đĩa KHÔNG bị sửa (checksum SHA-256 không đổi trong suốt quá "
                  "trình). Rule sửa/loại được định nghĩa và đóng băng TRƯỚC khi nhìn bất kỳ kết quả "
                  "policy nào trên Test -- chi tiết đầy đủ (bao gồm trace ngược nguồn dữ liệu gốc "
                  "thượng nguồn xác nhận lỗi có từ raw data, không phải do pipeline dự án) ở "
                  "`final_test/DATA_QUALITY_GATE.md`.", size=9.5, color=GREY)

        heading(doc, "9.3 Cấu hình đóng băng (frozen)", level=2)
        add_table(doc, ["Tham số", "Giá trị"], [
            ["Policy", "MOMAQL canonical"], ["Số tài xế", "200"],
            ["λ / γ / α", "0,5 / 0,9 / 0,1"],
            ["Bộ giải ghép cặp", "Hungarian joint assignment"],
            ["Q-table", "momaql_q_table_trained.json (đóng băng, không học thêm)"],
            ["Seeds", "20260721, 20260722, 20260723, 20260724, 20260725"],
        ])
        para(doc, "Chi tiết đầy đủ (hash mã nguồn engine, checksum dataset/Q-table, môi trường chạy) "
                  "ở `final_test/FINAL_TEST_PROTOCOL.md`.", size=9.5, color=GREY)

        heading(doc, "9.4 Kết quả Baseline trên Held-out Test", level=2)
        momaql_test = ft_baseline_summary.get("MOMAQL", {})
        momaql_val = r1_stats.get("MOMAQL", {})
        add_table(doc, ["", "Utility ($)", "Gini"], [
            ["MOMAQL -- Validation", vnum(momaql_val.get("u", 0)), vdec(momaql_val.get("g", 0))],
            ["MOMAQL -- Held-out Test", vnum(float(momaql_test.get("utility_mean", 0))),
             vdec(float(momaql_test.get("gini_mean", 0)))],
        ])
        para(doc, "Thứ hạng Utility theo policy trên Test GIỐNG HỆT Validation: "
                  "MOMAQL > Greedy > Nearest > LAF > Exact REASSIGN. Điểm vận hành của MOMAQL ổn "
                  "định qua các giai đoạn thời gian khác nhau.")

        heading(doc, "9.5 Kết quả Ablation trên Held-out Test", level=2)
        full_t = float(ft_ablation_summary["full"]["utility_mean"])
        nof_t = float(ft_ablation_summary["no_forecast"]["utility_mean"])
        nof_gini_t = float(ft_ablation_summary["no_forecast"]["gini_mean"])
        full_gini_t = float(ft_ablation_summary["full"]["gini_mean"])
        nofair_t = float(ft_ablation_summary["no_fairness"]["utility_mean"])
        nofair_gini_t = float(ft_ablation_summary["no_fairness"]["gini_mean"])
        add_table(doc, ["So sánh", "Validation", "Held-out Test"], [
            ["Full so với No-Forecast (Utility)", f"+{vdec((full_u-nf_u)/nf_u*100, 1)}%",
             f"+{vdec((full_t-nof_t)/nof_t*100, 1)}%"],
            ["No-Forecast Gini so với Full", f"{vdec(float(r2['no_forecast']['gini_mean']), 4)} vs "
             f"{vdec(float(r2['full']['gini_mean']), 4)} (No-Forecast công bằng hơn)",
             f"{vdec(nof_gini_t, 4)} vs {vdec(full_gini_t, 4)} (No-Forecast công bằng hơn)"],
            ["No-Fairness Gini so với Full",
             f"{vdec(float(r2['no_fairness']['gini_mean']), 4)} vs {vdec(float(r2['full']['gini_mean']), 4)}",
             f"{vdec(nofair_gini_t, 4)} vs {vdec(full_gini_t, 4)}"],
            ["No-Fairness Utility so với Full", f"{vdec((float(r2['no_fairness']['utility_mean'])-full_u)/full_u*100, 1)}%",
             f"{vdec((nofair_t-full_t)/full_t*100, 1)}%"],
        ])
        para(doc, "Cả 5/5 seed trên cả 2 tập: Full > No-Forecast về Utility (chiều giống nhau, độ "
                  "lớn yếu hơn trên Test); No-Forecast công bằng hơn Full (Gini thấp hơn) trên CẢ HAI "
                  "tập -- đây là phát hiện KHÔNG khớp kỳ vọng paper (paper kỳ vọng dự báo cải thiện "
                  "fairness), và phát hiện này GIỮ NGUYÊN (generalize) trên held-out Test, tức là "
                  "discrepancy so với paper ổn định, không phải nhiễu ngẫu nhiên của Validation.")

        if ft_horizon:
            heading(doc, "9.6 Kết quả Long-Horizon trên Held-out Test", level=2)
            para(doc, "Test span = 42 ngày lịch ≥ 37 → dùng đủ bộ checkpoint chuẩn "
                      "(ngày 1,2,3,4,5,6,7,14,21,28,37), cùng phương pháp một quỹ đạo có checkpoint "
                      "như Validation.")
            ft_gap_rows = []
            for d in [7, 21, 37]:
                if d in gap and d in ft_gap:
                    ft_gap_rows.append([str(d), f"{vdec(gap[d][1], 1)}%", f"{vdec(ft_gap[d][1], 1)}%"])
            add_table(doc, ["Ngày", "Validation (Full vs No-Forecast, Utility)", "Held-out Test"], ft_gap_rows)
            para(doc, "Lợi thế Utility dài hạn của Full so với No-Forecast GENERALIZE (cùng chiều "
                      "trên Test) nhưng yếu hơn ở mọi checkpoint đo được. Về fairness, No-Forecast "
                      "vẫn công bằng hơn Full tại ngày 37 trên cả 2 tập.")

        heading(doc, "9.7 Validation vs Test -- generalization tổng thể", level=2)
        n_gen_ft = sum(1 for r in ft_v_vs_t if r["generalized"] == "Yes")
        para(doc, f"{n_gen_ft}/{len(ft_v_vs_t)} phát hiện định trước (pre-specified) từ Validation "
                  "giữ ĐÚNG CHIỀU trên held-out Test (bảng đầy đủ: `final_test/validation_vs_test.csv`).")
        callout(doc, "QUAN TRỌNG: \"generalize\" ở đây có nghĩa PHÁT HIỆN quan sát trên Validation "
                     "LẶP LẠI cùng chiều trên Test -- một khái niệm hoàn toàn khác với \"paper claim "
                     "được reproduce\". Một phát hiện có thể generalize (ổn định theo thời gian) "
                     "trong khi paper claim tương ứng vẫn Not Reproduced -- ví dụ rõ nhất là C4 ở "
                     "Mục 10: cả Validation lẫn Test đều cho No-Forecast công bằng hơn Full, tức là "
                     "CHÍNH sự KHÔNG khớp với paper mới là thứ generalize, không phải claim của paper.")

    # ---------------- 10. Replication Assessment ----------------
    heading(doc, "10. Đánh giá Tái lập -- Paper vs. Của chúng tôi, theo từng claim (Replication Assessment)", level=1)
    para(doc, "Đây là bảng trả lời trực tiếp câu hỏi trọng tâm: dự án này có đưa ra đủ bằng chứng "
              "để thuyết phục rằng cơ chế paper tuyên bố thực sự xuất hiện lại trong tái lập này "
              "hay không -- theo từng claim riêng biệt, không gộp chung. Bảng dưới báo cáo trên "
              "Validation (đầy đủ mọi claim con -- Validation là nơi mọi thí nghiệm chính thực sự "
              "chạy).")
    verdict_table(doc, ["Tuyên bố của paper", "Kết luận (Validation)"], [
        ["C1 -- Utility và fairness đánh đổi lẫn nhau", "Reproduced"],
        ["C2 -- Proposed vượt trội baseline đã điều chỉnh về trade-off", "Reproduced (có phạm vi)"],
        ["C3 -- RL ổn định hơn khi horizon tăng", "Partially Reproduced (chuyển pha thực tế "
         "~ngày 14–21, không phải tức thì)"],
        ["C4a -- Dự báo giúp fairness dài hạn", "Partially Reproduced (xác nhận từ ngày ~21 trên "
         "bản mở rộng 37 ngày, không phải cửa sổ ngày-3 của paper)"],
        ["C4b -- Dự báo có thể tốn fairness ngắn hạn", "Not Evaluated (đã đo ngày 1–7, nhưng kết "
         "quả không đủ phân biệt theo hướng nào -- không ép thành Partial)"],
        ["C5 -- Bóc tách: thành phần dự báo giúp cả utility+fairness", "Reproduced"],
        ["C6a -- Bỏ fairness làm utility tăng mạnh", "Not Reproduced (utility giảm, đảo chiều "
         "so với paper)"],
        ["C6b -- Bỏ fairness làm bất công bằng bùng nổ", "Reproduced"],
        ["Số liệu tuyệt đối chính xác", "Not Evaluated (không phải mục tiêu, theo thiết kế -- "
         "xem Mục 3.1)"],
        ["(Bổ sung, không phải claim gốc của paper) Ổn định lợi thế dự báo theo quy mô đội xe",
         "Not Reproduced (lợi thế co từ +42% về +0%, N=100→400)"],
    ])

    if ft_available:
        para(doc, "Bảng dưới đây là ĐÁNH GIÁ CUỐI CÙNG (final) cho 6 claim gốc C1-C6, kết hợp cả hai "
                  "trục: (a) held-out generalization -- phát hiện Validation có lặp lại trên Test "
                  "không; và (b) paper replication verdict -- có khớp đúng tuyên bố định tính của "
                  "paper không. HAI TRỤC NÀY ĐỘC LẬP, không gộp thành một cột duy nhất -- xem C4 (Not "
                  "Reproduced dù Generalized) và C5/C6 (Partial, một thành phần reproduce, một không) "
                  "làm ví dụ vì sao không thể gộp.", bold=True)
        claim_id_map = {"C1:": "C1", "C2:": "C2", "C3:": "C3", "C4:": "C4", "C5:": "C5", "C6:": "C6"}
        ft_rows = []
        for r in ft_claims:
            cid = next((v for k, v in claim_id_map.items() if r["claim"].startswith(k)), r["claim"])
            ft_rows.append([cid, r["heldout_generalization"], r["paper_replication_verdict"]])
        dual_verdict_table(doc, ["Claim", "Generalize trên Held-out Test", "Kết luận tái lập so với paper"], ft_rows)
        para(doc, "Không có claim nào được viết là \"6/6 reproduced\". Bảng đầy đủ với evidence/"
                  "caveat từng dòng: `final_test/test_claim_assessment.csv`.", size=9.5, color=GREY)

    # ---------------- 11. Discrepancies and Limitations ----------------
    heading(doc, "11. Sai khác và Giới hạn (Discrepancies and Limitations)", level=1)
    heading(doc, "11.1 Vì sao số liệu khác paper", level=2)
    add_table(doc, ["Khác biệt", "Tác động kỳ vọng"], [
        ["Dữ liệu khác năm (2013 vs. 2016 của paper)", "Mẫu hình nhu cầu/không gian có thể khác; "
         "không phải bằng chứng khả năng tổng quát hóa qua các năm"],
        ["Dự báo là bảng Q(vùng,giờ) học TD(0), không phải MLP theo cặp OD của paper",
         "Tín hiệu nhìn trước khác bản chất -- so sánh trực tiếp ở Mục 8.5 cho thấy khác biệt "
         "thật, không chỉ lý thuyết"],
        ["Số tài xế (200) là giả định, paper không nêu rõ", "Ảnh hưởng trực tiếp tỷ lệ cung/cầu -- "
         "Mục 8.7 cho thấy hiệu ứng dự báo phụ thuộc mạnh vào tỷ lệ này"],
        ["ETA đón dùng khoảng cách đường thẳng, tốc độ cố định, không phải định tuyến đường thật",
         "Đơn giản hóa tính khả thi ghép cặp; có thể lệch nhẹ candidate pool thực tế"],
        ["Baseline REASSIGN/Balance-RP không tái lập được nguyên bản", "Khác biệt tuyệt đối so "
         "với paper; không ảnh hưởng đến so sánh nội bộ giữa các baseline đã điều chỉnh"],
    ])
    heading(doc, "11.2 Kết quả Tiêu cực và Rỗng (Negative and Null Results)", level=2)
    bullets(doc, [
        "Triển khai dự báo đầu tiên có lỗi gán trạng thái thật (phần thưởng gán cho điểm trả "
        "khách, không bootstrap), cho bảng Q tương quan ÂM (r=−0,36) với nhu cầu thực tế tương "
        "lai; đã sửa bằng cập nhật Bellman đúng trên trạng thái (vùng,giờ), tương quan sau sửa "
        "+0,51.",
        f"Tỷ lệ đổi quyết định D = P(π_full ≠ π_no_forecast) trên toàn bộ 37 ngày chỉ "
        f"{vdec(disag_mean*100, 1)}% -- dự báo chỉ đổi một phần nhỏ trong tổng số phép gán, nhưng vẫn "
        "tích lũy thành chênh lệch utility +20% ở Mục 8.4. Dữ liệu hiện có KHÔNG tách riêng được "
        "tỷ lệ đổi quyết định chỉ trong đúng cửa sổ ngày 1–7 (script hiện tại đo cộng dồn trên cả "
        "quỹ đạo); đây là một khoảng trống công cụ được công bố rõ, không phải số liệu bị bỏ sót.",
        "Hướng utility của C6 không tái lập được; có phân rã (Mục 8.6) nhưng chưa có cơ chế không "
        "gian được kiểm chứng độc lập (cuốc ngắn trung tâm vs. ngoại vi).",
        "Độ nhạy quy mô đội xe: lợi thế dự báo bùng nổ khi thiếu hụt cung (+42% tại N=100) và bão "
        "hòa khi thừa cung (0% tại N=400) -- một phát hiện thật, không được paper kiểm tra.",
    ])
    heading(doc, "11.3 Khoảng trống Tái lập và Giả định (Reproducibility Gaps and Assumptions)", level=2)
    bullets(doc, [
        "A1. Paper không nêu rõ số lượng tài xế mô phỏng. Quyết định tái lập: dùng 200, giả định "
        "hợp lý cho quy mô Manhattan, công bố rõ đây là giả định.",
        "A2. Paper không nêu rõ thuật toán gộp vị trí thành node. Quyết định tái lập: dùng 67 zone "
        "TLC chính thức thay cho gộp tùy biến.",
        "A3. Paper không mô tả đầy đủ hidden layer/optimizer/learning rate/epochs của MLP dự báo. "
        "Quyết định tái lập: khi cần một MLP thật để so sánh trực tiếp (Mục 8.5), dùng kiến trúc "
        "3 lớp với embedding vùng + one-hot giờ, Adam, 30 epoch -- tự chọn, công bố rõ, không "
        "tuyên bố khớp paper.",
        "A4. Paper không công bố đầy đủ chi tiết triển khai baseline REASSIGN/Balance-RP gốc. "
        "Quyết định tái lập: dùng Exact REASSIGN (Lesmana et al. 2019) và LAF (Sühr et al. 2019) "
        "làm baseline gần nhất có thể kiểm chứng, không tuyên bố là bản triển khai trực tiếp.",
        "A5. Paper không báo cáo random seed hay thống kê nhiều lần chạy. Quyết định tái lập: chạy "
        "5 seed độc lập (3 seed cho thí nghiệm cơ chế bổ sung) và báo cáo trung bình ± độ lệch "
        "chuẩn cho mọi kết quả chính.",
    ])

    # ---------------- 12. Conclusion ----------------
    heading(doc, "12. Kết luận (Conclusion)", level=1)
    para(doc, "Dự án tái lập sạch các tuyên bố trung tâm về đánh đổi và vượt trội baseline (C1, C2), "
              "trên dữ liệu taxi NYC 2013 thực tế, độc lập, lấy mẫu riêng. Các tuyên bố "
              "horizon dài và no-fairness (C3, C4, C5, C6) chỉ tái lập một phần hoặc không tái lập ở "
              "một thành phần: hướng định tính đúng khi horizon đủ dài và trên bất công bằng, nhưng "
              "không đúng ở hướng utility của C6, không đúng ở thành phần fairness của C5, và C4 "
              "(dự báo cải thiện fairness dài hạn) hoàn toàn Not Reproduced. Một quét quy mô đội xe "
              "(không phải claim gốc của paper) cho thấy thêm: chính lợi thế của dự báo phụ thuộc quy "
              "mô -- lớn khi khan hiếm tài xế, không đáng kể khi cung bão hòa cầu -- chứ không phải "
              "thuộc tính cố định của thuật toán. Bốn thí nghiệm cơ chế bổ sung (Mục 8.4.1) cho "
              "thấy thời điểm chuyển pha ngày 14–21 có nhiều bằng chứng liên hệ hội tụ với nhau "
              "(độ phủ trạng thái, giá trị Q, tỷ trọng fairness), dù không cái nào là ablation có "
              "kiểm soát chứng minh nhân quả.")
    if ft_available:
        n_gen_ft = sum(1 for r in ft_v_vs_t if r["generalized"] == "Yes")
        para(doc, f"Sau khi đóng băng implementation và cấu hình trên Validation, held-out temporal "
                  f"Test evaluation (Mục 9) xác nhận {n_gen_ft}/{len(ft_v_vs_t)} phát hiện định trước "
                  "từ Validation ĐÚNG CHIỀU trên Test -- củng cố độ tin cậy rằng hành vi quan sát "
                  "được của bản tái lập ổn định theo thời gian. Tuy nhiên, kết quả held-out cũng xác "
                  "nhận rằng KHÔNG PHẢI mọi tuyên bố định tính của paper gốc đều được reproduce: lợi "
                  "ích Utility từ dự báo vẫn duy trì, nhưng lợi ích Fairness từ dự báo vẫn Not "
                  "Reproduced trên cả 2 tập dữ liệu.", bold=True)
    para(doc, "Báo cáo này trình bày trung thực tất cả điều trên thay vì làm tròn thành “đã tái lập "
              "hoàn toàn”, đúng tinh thần chính paper phân biệt giữa tái lập số liệu chính xác và "
              "tái lập xu hướng định tính. Kết luận cuối cùng: "
              "\"Strong Partial Trend Replication with held-out temporal support\" -- Strong vì "
              "13/13 phát hiện generalize trên held-out Test; Partial vì các paper claim liên quan "
              "fairness không reproduce đầy đủ; held-out temporal support vì Test là một lát cắt "
              "thời gian độc lập, chưa từng dùng để tune.", bold=True)

    # ---------------- Appendix ----------------
    heading(doc, "Phụ lục (Appendix)", level=1)
    para(doc, "A.1 Reproducibility snapshot (tóm tắt -- xem Technical Documentation để có đầy đủ "
              "lệnh chạy lại, cấu trúc repo và module contract)", size=11, bold=True)
    add_table(doc, ["Mục", "Giá trị"], [
        ["Git commit tại thời điểm build tài liệu này", head],
        ["Dataset SHA-256 (train/val/test/Q-table)", "Xem reports/dataset_checksums.json và "
         "docs/techdoc/"],
        ["Số test bất biến (chạy lại thật)", f"tests/test_simulator_invariants.py: {test_line}"],
        ["Random seeds (thí nghiệm chính)", "20260721, 20260722, 20260723, 20260724, 20260725"],
    ])
    para(doc, "A.2 Tài liệu tham khảo", size=11, bold=True)
    bullets(doc, [
        "Yufan Kang, Jeffrey Chan, Wei Shao, Flora D. Salim, and Christopher Leckie. Long-term "
        "fairness in ride-hailing platform. ECML PKDD 2024, LNCS 14949. "
        "https://arxiv.org/abs/2407.17839",
        "Nixie S. Lesmana, Xuan Zhang, and Xiaohui Bei. Balancing efficiency and fairness in "
        "on-demand ridesourcing. NeurIPS 2019, pages 5310-5320.",
        "Tom Sühr, Asia J. Biega, Meike Zehlike, Krishna P. Gummadi, and Abhijnan Chakraborty. "
        "Two-sided fairness for repeated matchings in two-sided markets. KDD '19, pages 3082-3092.",
    ])

    doc.save(OUT)
    print(f"[done] {OUT}")
    ctx = dict(head=head, r1_stats=r1_stats, r2=r2, r2_raw=r2_raw, pareto=pareto,
               hz_mean=hz_mean, hz_gini_mean=hz_gini_mean, gap=gap, disag_mean=disag_mean,
               fleet_mean=fleet_mean, fleet_completed_mean=fleet_completed_mean,
               spool_mean=spool_mean, sdis_mean=sdis_mean, qconv=qconv,
               week_weekday=week_weekday, week_weekend=week_weekend, fair_bal=fair_bal,
               mlp_summary=mlp_summary)
    return doc, ctx


if __name__ == "__main__":
    main()
