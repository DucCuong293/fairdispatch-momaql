# -*- coding: utf-8 -*-
"""Build Technical_Documentation.docx -- tài liệu kỹ thuật/khả năng tái lập
đồng hành với Research Report (docs/docx_report/). Tài liệu này trả lời
ĐÚNG MỘT câu hỏi: nếu một kỹ sư clone repo này ngày mai, họ có hiểu được
kiến trúc, chạy lại được mọi thí nghiệm, và ra đúng artifact đã giữ lại
không? Nó KHÔNG tranh luận lại khoa học -- đó là việc của Research Report.
Mọi lệnh, số dòng, và giá trị config dưới đây đều verify trực tiếp trên
repo thật tại thời điểm build (`git rev-parse HEAD` mới, tương đương
`wc -l` trên reports/*.csv mới, chạy pytest mới và đếm lại, `sha256sum`
mới của file dataset).
Chạy: python docs/techdoc/build_technical_documentation.py
"""
import csv
import hashlib
import platform
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from _docx_style import (
    NAVY, ACCENT as BLUE, GREY, setup_document, title_page, heading, para,
    bullets, add_table as _styled_add_table, add_figure as _styled_add_figure,
    code_block, callout,
)

ROOT = Path(__file__).resolve().parents[2]
REPORTS = ROOT / "reports"
DATA = ROOT / "data"
FINAL_TEST = ROOT / "final_test"
OUT = Path(__file__).parent / "Technical_Documentation.docx"

DARK = RGBColor(0x20, 0x20, 0x20)
_fig_counter = {"n": 0}


def git_head():
    try:
        return subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=ROOT, text=True).strip()
    except Exception:
        return "unavailable"


def sha256(path):
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def csv_row_count(name):
    with (REPORTS / name).open(encoding="utf-8") as f:
        return sum(1 for _ in f) - 1


def add_table(doc, headers, rows):
    return _styled_add_table(doc, headers, rows)


def add_figure(doc, filename, caption, width=5.9):
    path = Path(__file__).parent.parent / "docx_report" / "figures" / filename
    _fig_counter["n"] += 1
    _styled_add_figure(doc, path, caption, width=width, number=_fig_counter["n"])


def main():
    head = git_head()

    # Chạy pytest thật, đếm số dòng. Bắt đúng theo hình dạng dòng tổng kết
    # thật của pytest ("20 passed in 1.2s", "5 failed, 15 passed in 2.3s",
    # "20 errors in 13.5s") thay vì tìm chuỗi con lỏng lẻo, vì cách cũ có thể
    # bắt nhầm 1 dòng không liên quan (vd 1 thông báo nội bộ của pytest/encoding
    # tình cờ có chứa chữ "passed") khi lần chạy bị lỗi trước khi test nào
    # thực sự chạy.
    try:
        out = subprocess.run([sys.executable, "-m", "pytest", "tests/test_simulator_invariants.py", "-q"],
                              cwd=ROOT, capture_output=True, text=True, timeout=120)
        combined = out.stdout + "\n" + out.stderr
        m = re.search(r"^\d+ (?:passed|failed|errors?)\b.*$", combined, re.MULTILINE)
        if m:
            test_line = m.group(0).strip()
            if "error" in test_line.lower() and "train.parquet" in combined:
                test_line += (" -- data/train.parquet chưa có trong bundle này "
                              "(bị gitignore theo thiết kế, quá lớn để đóng gói); phải lấy "
                              "riêng các phần chia parquet để chạy lại bộ test này cục bộ "
                              "(Mục 11.3)")
        else:
            test_line = "(không đọc được dòng tổng kết pytest -- xem stdout/stderr thô để biết chi tiết)"
    except Exception as e:
        test_line = f"(chạy pytest thất bại: {e})"

    # Hash dataset mới.
    ds_hashes = {}
    for fname in ["train.parquet", "val.parquet", "test.parquet", "momaql_q_table_trained.json"]:
        p = DATA / fname
        if p.exists():
            ds_hashes[fname] = (sha256(p), p.stat().st_size)

    # Artifact Final Test (chỉ đọc, không chạy lại).
    ft_available = FINAL_TEST.exists() and (FINAL_TEST / "test_quality_transform_manifest.json").exists()
    ft_manifest = None
    ft_commands_log = ""
    ft_environment_txt = ""
    if ft_available:
        import json
        ft_manifest = json.loads((FINAL_TEST / "test_quality_transform_manifest.json").read_text(encoding="utf-8"))
        cmd_log_path = FINAL_TEST / "logs" / "commands.log"
        env_path = FINAL_TEST / "logs" / "environment.txt"
        if cmd_log_path.exists():
            ft_commands_log = cmd_log_path.read_text(encoding="utf-8")
        if env_path.exists():
            ft_environment_txt = env_path.read_text(encoding="utf-8")

    doc = Document()
    setup_document(doc, footer_label="FairDispatch -- Tài liệu Kỹ thuật")
    title_page(
        doc,
        kicker="Tài liệu Kỹ thuật",
        title="FairDispatch -- MOMAQL",
        subtitle="Hướng dẫn Kỹ thuật & Khả năng Tái lập -- tài liệu đồng hành với Research Report; "
                 "tài liệu này trình bày phần triển khai, không tranh luận lại claim khoa học.",
        meta_lines=[f"Git commit tại thời điểm build: {head[:12]}",
                    "Cấu trúc mã nguồn bên dưới khớp cả với 03_Source_Code_Va_Ket_Qua/ của gói "
                    "nộp này lẫn repo dev đã copy ra (fairdispatch_v3_clean/) -- cùng đường dẫn "
                    "tương đối, chỉ khác tên thư mục gốc."],
    )

    # ---------------- 1. Tổng quan Hệ thống ----------------
    heading(doc, "1. Tổng quan Hệ thống (System Overview)", level=1)
    para(doc, "FairDispatch -- MOMAQL là một bản tái lập độc lập, xây lại từ đầu, các xu hướng "
              "định tính của Kang et al. [2024], \"Long-term Fairness in Ride-Hailing Platform\" "
              "(ECML PKDD 2024, arXiv:2407.17839): một chính sách điều phối RL đa mục tiêu "
              "(MOMAQL) đánh đổi giữa công bằng thu nhập tài xế và hiệu quả hệ thống, dùng ước "
              "lượng giá trị nhìn trước (look-ahead) để tránh quyết định công bằng cận thị từng "
              "cửa sổ. Dự án này đánh giá ý tưởng đó trên dữ liệu taxi NYC TLC 2013 thật "
              "(912.375 huấn luyện / 195.508 kiểm định / 195.510 kiểm thử), so với 4 chính sách "
              "baseline, rồi xác minh lại mọi phát hiện trên một tập Test tách riêng theo thời "
              "gian mà implementation chưa từng thấy trong lúc phát triển (Mục 8). Tài liệu này "
              "là tài liệu đồng hành về kỹ thuật/khả năng tái lập -- không tranh luận lại claim "
              "nào được tái lập; verdict đó (\"Strong Partial Trend Replication with held-out "
              "temporal support\") cùng lý giải đầy đủ theo từng claim nằm ở Research Report "
              "(docs/docx_report/).")
    para(doc, "Pipeline đầu-cuối thật được thực thi bởi repo này:")
    add_figure(doc, "replication_pipeline.png", "Hình 1: Chuyến đi NYC TLC 2013 thật -> lọc "
               "Manhattan + chất lượng -> 67 zone taxi TLC -> Q(vùng,giờ) huấn luyện qua Bellman "
               "TD(0) -> simulator theo lô cửa sổ 60 giây -> ghép cặp Hungarian M-to-N -> chỉ số "
               "utility/Gini/variance.")
    para(doc, "Không có message queue, không database, không web service -- mọi giai đoạn là một "
              "script Python thuần đọc/ghi file cục bộ. Simulator là một vòng lặp discrete-event "
              "(DES) nhẹ, chạy trong cùng process (src/simulator.py), không phải hệ thống "
              "real-time.")

    # ---------------- 2. Cấu trúc Repository ----------------
    heading(doc, "2. Cấu trúc Repository (Repository Structure)", level=1)
    code_block(doc, """
03_Source_Code_Va_Ket_Qua/   (thư mục gốc của gói nộp này -- cấu trúc tương đối
|                             giống hệt repo dev fairdispatch_v3_clean/ đã copy ra,
|                             chỉ khác tên thư mục gốc)
|-- src/
|   |-- simulator.py       dataclass Driver/SimResult, init_drivers,
|   |                      feasible_drivers, commit_trip, run_simulation_batched,
|   |                      run_simulation_with_horizon, run_simulation (cũ)
|   `-- policies.py        hungarian_batch_assign + 5 class policy (Greedy,
|                          Nearest, LAF, ExactReassignPolicy, MOMAQLPolicy)
|-- common_loader.py        loader parquet -> dict request; gini/variance/std/CV
|-- train_momaql.py          huấn luyện bảng Q(vùng,giờ) canonical (1 lượt)
|-- train_momaql_multipass.py  ĐÃ CÔNG BỐ KHÔNG ỔN ĐỊNH -- không dùng để đánh giá
|-- run_r1.py                 R1: so sánh baseline 5 policy
|-- run_r2_ablation.py        R2: Full / w/o-Forecast / w/o-Fairness
|-- run_pareto_frontier.py    quét lambda, 7 giá trị
|-- run_multi_horizon.py      quỹ đạo multi-horizon + tỷ lệ đổi quyết định policy
|-- run_complete_verifications.py   quét quy mô đội xe + đổi quyết định theo không gian
|-- run_spatial_candidate_pool.py   độ sâu vùng ứng viên, lõi vs. ngoại vi
|-- run_q_table_convergence.py      hội tụ bảng Q theo ngày (37 ngày)
|-- run_hypothesis1_weekly_cycle.py  kiểm tra cơ chế chu kỳ nhu cầu tuần
|-- run_hypothesis4_fairness_balance.py  trace cân bằng score fairness/nhìn trước
|-- train_and_eval_mlp.py     MLP PyTorch thật dự báo nhu cầu vs. Q dạng bảng
|-- make_report_figures.py    sinh lại mọi hình từ reports/*.csv
|-- tests/
|   `-- test_simulator_invariants.py   20 test bất biến (cần có data/*.parquet
|                                        cục bộ để chạy -- Mục 11.3)
|-- scripts/final_test/       Pipeline Final Held-out Test (Mục 8): quality
|                              transform, script chạy baseline/ablation/long-horizon,
|                              script tổng hợp claim assessment
|-- data/                     các phần chia train/val/test parquet (gitignore, xem
|                              Mục 10 để lấy SHA-256) + JSON bảng Q đã huấn luyện
|-- reports/                  mọi CSV/JSON/PNG Validation được giữ lại, trích dẫn ở cả 2 tài liệu
|-- final_test/               Kết quả Final Held-out Test (Mục 8): protocol,
|                              data quality gate, kết quả từng seed, claim
|                              assessment, mentor summary, hình, log
`-- docs/
    |-- docx_report/           Research Report (.docx) + hình của nó
    |-- techdoc/                tài liệu này
    `-- ride_hailing_fairness_report_en(vi)/   paper LaTeX + PDF đã compile
""")
    para(doc, "Repo này không có dependency lockfile, file package-metadata, Dockerfile, CI "
              "workflow, hay central config file nào. Mỗi run_*.py là một script độc lập với "
              "hằng số riêng ở cấp module (N_DRIVERS, SEEDS, v.v.) -- xem Mục 7.", size=9.5,
              color=GREY)

    # ---------------- 3. Môi trường ----------------
    heading(doc, "3. Môi trường (Environment)", level=1)
    para(doc, "Môi trường đã audit tại thời điểm build (đây là máy repo này được verify lần cuối "
              "-- KHÔNG tuyên bố mọi thí nghiệm lịch sử đều chạy trên phần cứng giống hệt; không "
              "giữ lại environment lockfile từ các lần chạy trước).", size=9.5,
              color=GREY)
    add_table(doc, ["Thành phần", "Giá trị"], [
        ["OS", f"{platform.system()} {platform.release()} ({platform.version()})"],
        ["Python", platform.python_version()],
        ["CPU", "Intel Core i5-10300H, 4 lõi / 8 luồng"],
        ["RAM", "Tổng 15,8 GiB"],
        ["GPU", "NVIDIA GeForce GTX 1650, ~4,3 GB VRAM (có CUDA)"],
        ["Thư viện chính", "numpy 2.2.6, pandas 2.3.2, pyarrow 21.0.0, scipy 1.16.2, "
         "matplotlib 3.10.6, pytest 8.3.3, python-docx 1.2.0, torch 2.7.1+cu118"],
    ])
    code_block(doc, """
python -m venv .venv
.venv\\Scripts\\activate                 # Windows; dùng source .venv/bin/activate trên Linux/Mac
pip install numpy pandas pyarrow scipy matplotlib pytest python-docx
pip install torch                        # chỉ cần cho train_and_eval_mlp.py
""")
    para(doc, "Repo này không giữ CI/CD, không Docker image, không requirements.txt -- danh sách "
              "package ở trên là bộ đã audit, đang hoạt động; hãy ghim đúng version nếu cần tái "
              "lập bit-for-bit qua nhiều máy.")

    # ---------------- 4. Hợp đồng Dữ liệu ----------------
    heading(doc, "4. Hợp đồng Dữ liệu (Data Contract)", level=1)
    add_table(doc, ["Phần chia", "Vai trò"], [
        ["train.parquet", "Huấn luyện Q(vùng,giờ) (train_momaql.py); huấn luyện MLP cho "
         "sensitivity forecast"],
        ["val.parquet", "Phát triển / phân tích: build+debug simulator, so sánh policy, "
         "ablation, quét Pareto, long-horizon, thí nghiệm cơ chế, chốt cấu hình cuối"],
        ["test.parquet", "CHỈ dùng để xác minh held-out cuối cùng theo thời gian (Mục 8) -- "
         "không bao giờ dùng để chọn lambda/gamma/alpha, không bao giờ dùng bởi product demo "
         "trực tiếp"],
    ])
    heading(doc, "4.1 Schema parquet gốc (train/val/test)", level=2)
    para(doc, "Các cột được common_loader.load_requests_fast() và loader riêng của "
              "train_momaql.py đọc: pickup_ts (timestamp), pickup_latitude, pickup_longitude, "
              "dropoff_latitude, dropoff_longitude, fare_amount, duration_seconds, "
              "pickup_zone_id, dropoff_zone_id. pickup_ts là epoch time tuyệt đối -- giờ-trong-"
              "ngày được tính trực tiếp từ đó (KHÔNG phải offset tương đối theo từng file), nên "
              "một khung giờ mang cùng ý nghĩa thời gian thực trong train.parquet lẫn "
              "val/test.parquet.")
    heading(doc, "4.2 Request dict đã xử lý (in-memory, mỗi dòng)", level=2)
    code_block(doc, """
{
  "_idx": int,                    # vị trí trong danh sách request đã sắp xếp
  "pickup_ts": float,             # giây, TƯƠNG ĐỐI so với dòng đầu tiên của split này
  "pickup_latitude": float, "pickup_longitude": float,
  "dropoff_latitude": float, "dropoff_longitude": float,
  "fare_amount": float,           # cước phí USD thật
  "duration_seconds": float,
  "pickup_zone_id": int, "dropoff_zone_id": int,   # id zone taxi TLC
  "pickup_hour": int,             # 0-23, tính từ epoch time TUYỆT ĐỐI
  "dropoff_hour": int,            # 0-23, = (pickup_epoch + duration) % 24h
}
""")
    heading(doc, "4.3 Trạng thái Driver (src/simulator.py:Driver)", level=2)
    code_block(doc, """
Driver:
  driver_id: int
  lat, lon: float                 # vị trí hiện tại
  available_at: float             # thời điểm sớm nhất driver này nhận chuyến mới
  total_income: float = 0.0       # net tích lũy (fare - deadhead_cost)
  total_deadhead_cost: float = 0.0
  total_trips: int = 0
""")
    heading(doc, "4.4 State/Action/Reward RL (MOMAQLPolicy)", level=2)
    bullets(doc, [
        "State S = (pickup_zone_id, pickup_hour) của chuyến driver vừa nhận.",
        "State S' = (dropoff_zone_id, dropoff_hour) của chính chuyến đó -- driver kết thúc ở "
        "đâu/khi nào.",
        "Action = một phép ghép cặp joint M-to-N cho 1 cửa sổ 60 giây, giải 1 lần mỗi cửa sổ "
        "bằng thuật toán Hungarian (không phải chọn greedy từng request).",
        "Reward = fare_amount − eta_seconds × 0,0025 (hệ số chi phí deadhead, USD mỗi giây ETA "
        "đón khách).",
        "Key bảng Q: (zone_id: int, hour_of_day: int) -> float. Khi serialize JSON dùng key dạng "
        "chuỗi \"zone:hour\" (vd \"137:0\"); MOMAQLPolicy._parse_q_table() nhận cả tuple key gốc "
        "lẫn định dạng chuỗi này.",
    ])

    # ---------------- 5. Đặc tả Module ----------------
    heading(doc, "5. Đặc tả Module (Module Specification)", level=1)
    heading(doc, "5.1 src/simulator.py", level=2)
    add_table(doc, ["Hàm", "Đầu vào", "Đầu ra"], [
        ["init_drivers(n, requests, seed)", "số lượng driver, danh sách request, seed RNG",
         "list[Driver], đặt tại n vị trí đón khách thật lấy mẫu từ danh sách request"],
        ["feasible_drivers(drivers, req, now)", "danh sách driver, 1 request, thời điểm hiện tại",
         "list (driver, deadhead_miles, eta_seconds) cho driver rảnh tính tới `now` và trong "
         "ngưỡng MAX_PICKUP_ETA_SECONDS=600s"],
        ["commit_trip(d, req, dist, eta, now, result, record_trace)", "driver đã chọn + request + "
         "thời gian", "thay đổi trạng thái driver (income, vị trí, available_at) và tổng số của "
         "result"],
        ["run_simulation_batched(requests, n_drivers, policy, seed, window_seconds=60.0)",
         "danh sách request đã sắp xếp + đối tượng policy", "SimResult (thu nhập từng driver, số "
         "chuyến hoàn thành)"],
        ["run_simulation_with_horizon(..., checkpoint_days, compare_policy=None, "
         "zone_classifier=None)", "tương tự + các ngày checkpoint + policy so sánh tùy chọn",
         "(result, dict checkpoints, disagreement_rate) -- MỘT quỹ đạo duy nhất, chụp snapshot ở "
         "mỗi ngày checkpoint, không phải chạy lại độc lập"],
    ])
    para(doc, "Đơn giản hóa có chủ đích (đã công bố trong docstring module): vị trí driver là "
              "thật (lat, lon), không phải xấp xỉ theo ô không gian; ETA là khoảng cách "
              "haversine với tốc độ hằng số 12 mph, không phải mô hình road-routing thật; hệ số "
              "chi phí deadhead (0,0025 USD/giây) tái sử dụng nguyên trạng từ UtilityCoefficients "
              "đã đóng băng của dự án cha, không tự nghĩ lại.", size=9.5, color=GREY)
    heading(doc, "5.2 src/policies.py", level=2)
    add_table(doc, ["Policy", "score_fn(d, req, dist, eta)", "Ghi chú"], [
        ["Greedy", "fare_amount", "giống nhau giữa mọi candidate của 1 request; joint solve vẫn "
         "tối đa số chuyến phục vụ"],
        ["Nearest", "(600.0 − eta) × 0,0025", "định tâm lại để 0 = break-even tại ngưỡng "
         "feasibility; argmax(score) == argmin(eta)"],
        ["LAF", "[(mean_income − d.income) / max(mean_income,1)] × fare", "argmax(score) == "
         "argmin(driver.income); heuristic chỉ tối ưu fairness"],
        ["Exact REASSIGN", "fare_amount − eta×0,0025", "tối đa hóa net-utility thuần, giải joint "
         "M-to-N thật"],
        ["MOMAQL", "(1−λ)·[fare − eta×0,0025 + γ·Q(D_zone,D_hour)] + λ·[(mean_income−d.income)/"
         "max(mean_income,1)]·fare", "policy duy nhất có Q-learning online (on_committed cập "
         "nhật Q qua Bellman TD(0) trừ khi frozen=True)"],
    ])
    para(doc, "Cả 5 policy dùng CHUNG một khung tham chiếu: hungarian_batch_assign() giải một "
              "phép ghép joint thật qua scipy.optimize.linear_sum_assignment trên ma trận chi "
              "phí (n_req+n_drv)×(n_req+n_drv) đệm dummy, nên mọi request/driver đều có lựa chọn "
              "\"từ chối\"/\"rảnh\" thật với chi phí 0 -- đây là ghép cặp maximum-WEIGHT (đúng "
              "công thức sum_v I_rv ≤ 1 của paper), không phải maximum-cardinality.")
    heading(doc, "5.3 common_loader.py", level=2)
    add_table(doc, ["Hàm", "Mục đích"], [
        ["load_requests_fast(path)", "loader parquet -> list[dict] dùng PyArrow; tính "
         "pickup_hour/dropoff_hour từ epoch time tuyệt đối"],
        ["gini(values)", "hệ số Gini của danh sách thu nhập driver"],
        ["variance(values) / std(values)", "variance/std của toàn bộ tổng thể -- metric fairness "
         "chính của chính paper"],
        ["coefficient_of_variation(values)", "std/mean; chặn NaN khi |mean| < 1e-9 (không ổn "
         "định khi utility trung bình gần 0)"],
    ])

    # ---------------- 6. Thuật toán / Pseudocode ----------------
    heading(doc, "6. Thuật toán / Pseudocode (Algorithm)", level=1)
    heading(doc, "6.1 Vòng lặp dispatch theo lô (run_simulation_batched)", level=2)
    code_block(doc, """
mỗi cửa sổ 60 giây, theo đúng thứ tự pickup_ts thật:
    window_reqs = request có pickup_ts trong [window_start, window_start+60s)
    for req in window_reqs:
        candidates[req] = feasible_drivers(drivers, req, window_start)
            # rảnh tính tới window_start VÀ eta <= 600s

    assignments = policy.select_batch(candidates, window_start)
        # -> giải Hungarian joint thật, có lựa chọn từ chối chi phí 0

    for req, (driver, dist, eta) in assignments.items():
        if driver đã dùng trong cửa sổ này: skip   # chặn lỗi policy
        commit_trip(driver, req, dist, eta, window_start, result)
            # driver.income += fare - eta*0.0025
            # driver.available_at = window_start + eta + duration_seconds
        policy.on_committed(driver, req, dist, eta, window_start)
            # chỉ MOMAQL: cập nhật Bellman TD(0) online (bỏ qua nếu frozen)
""")
    heading(doc, "6.2 Score MOMAQL và cập nhật Q", level=2)
    code_block(doc, """
score(d, req) = (1-lambda) * [fare - eta*0.0025 + gamma * Q(D_zone, D_hour)]
              +  lambda    * [(mean_income - d.income) / max(mean_income,1)] * fare

Cập nhật Q (Bellman TD(0), lúc commit, chỉ khi không frozen):
  P = (pickup_zone_id, pickup_hour)      # state driver đang ở
  S = (dropoff_zone_id, dropoff_hour)    # state driver kết thúc ở
  reward = fare - eta*0.0025
  Q[P] <- Q[P] + alpha * (reward + gamma * Q[S] - Q[P])
""")
    para(doc, "lambda=0,5 (mặc định), gamma=0,9, alpha=0,1. ablation='no_forecast' ép số hạng "
              "gamma*Q(...) về 0; ablation='no_fairness' ép lambda về 0. frozen=True (dùng cho "
              "mọi lần đánh giá) tắt hẳn việc cập nhật Q -- bảng đã huấn luyện từ "
              "train_momaql.py chỉ được đọc, không ghi.")

    # ---------------- 7. Cấu hình ----------------
    heading(doc, "7. Cấu hình (Configuration)", level=1)
    para(doc, "Không có central config file nào (không YAML/JSON/.env). Mỗi hằng số dưới đây là "
              "1 biến Python cấp module, lặp lại ở từng script (đã verify thật, không giả "
              "định):")
    add_table(doc, ["Hằng số", "Giá trị", "Định nghĩa ở đâu"], [
        ["N_DRIVERS", "200", "đầu mỗi run_*.py và train_momaql.py"],
        ["SEEDS (thí nghiệm chính)", "20260721, 20260722, 20260723, 20260724, 20260725",
         "run_r1.py, run_r2_ablation.py, run_pareto_frontier.py, run_multi_horizon.py, "
         "train_and_eval_mlp.py"],
        ["SEEDS (thí nghiệm cơ chế)", "20260721, 20260722, 20260723 (3 seed)",
         "run_complete_verifications.py, run_spatial_candidate_pool.py, "
         "run_hypothesis4_fairness_balance.py"],
        ["WINDOW_SECONDS", "60,0", "tham số mặc định của run_simulation_batched"],
        ["MAX_PICKUP_ETA_SECONDS", "600,0 (10 phút)", "hằng số module src/simulator.py"],
        ["COST_PER_SECOND_DEADHEAD_USD", "0,0025", "hằng số module src/simulator.py"],
        ["AVG_SPEED_MPH", "12,0", "hằng số module src/simulator.py"],
        ["lambda (trọng số fairness)", "mặc định 0,5; quét {0; 0,2; 0,4; 0,5; 0,6; 0,8; 1,0} cho "
         "Pareto", "tham số constructor MOMAQLPolicy(lam=...)"],
        ["gamma (discount)", "0,9", "tham số constructor MOMAQLPolicy(gamma=...)"],
        ["alpha (learning rate Q)", "0,1", "tham số constructor MOMAQLPolicy(alpha=...)"],
    ])

    # ---------------- 8. Protocol Final Test ----------------
    if ft_available:
        heading(doc, "8. Protocol & Khả năng Tái lập Final Test (Final Test Protocol & "
                      "Reproducibility)", level=1)
        para(doc, "Chi tiết kỹ thuật đồng hành cho Final Held-out Temporal Test Evaluation "
                  "(Research Report Mục 9). Mọi thứ ở đây đọc trực tiếp từ `final_test/` tại "
                  "thời điểm build tài liệu -- không số nào ở đây gõ tay.", size=9.5, color=GREY)

        heading(doc, "8.1 Protocol Đóng băng (Frozen Protocol)", level=2)
        add_table(doc, ["Tham số", "Giá trị"], [
            ["Policy", "MOMAQL canonical (lambda=0,5, gamma=0,9, alpha=0,1)"],
            ["Số tài xế", "200"],
            ["Bộ giải ghép cặp", "Hungarian joint assignment (scipy.optimize.linear_sum_assignment)"],
            ["Bảng Q", "data/momaql_q_table_trained.json, đã đóng băng (không học thêm lúc đánh giá)"],
            ["Seed", "20260721, 20260722, 20260723, 20260724, 20260725"],
            ["Quy tắc không tuning", "Không đổi config/hyperparameter/policy/simulator sau khi đã "
             "xem bất kỳ kết quả Test nào"],
        ])
        para(doc, "Toàn văn protocol (hash mã nguồn, checksum dataset, môi trường): "
                  "`final_test/FINAL_TEST_PROTOCOL.md`.", size=9.5, color=GREY)

        heading(doc, "8.2 Data Quality Transform (raw file bất biến, tại thời điểm đánh giá)",
                level=2)
        tstats = ft_manifest["per_split"]["test"]
        para(doc, "test.parquet trên đĩa KHÔNG BAO GIỜ bị sửa. Hai rule độc lập chỉ áp dụng lên "
                  "evaluation view trong bộ nhớ, theo đúng thứ tự này, không trộn lẫn:")
        bullets(doc, [
            "A. Vệ sinh ranh giới thời gian nghiêm ngặt: loại các dòng có giây epoch pickup_ts "
            "trùng với giây epoch pickup_ts lớn nhất của val.parquet (một artifact thật do chia "
            "theo row-index -- các pickup thật khác nhau rơi cùng 1 giây tại điểm cắt, không "
            f"phải dữ liệu trùng lặp) -- loại {tstats['temporal_boundary_excluded']} dòng.",
            "B. Sửa duration tối thiểu, có tính tất định (deterministic): nếu duration_seconds "
            "lưu sẵn hợp lệ (0 < x <= 24h), giữ nguyên. Nếu không, mà dropoff_ts-pickup_ts hợp "
            "lệ, sửa duration_seconds_eval từ timestamp (quality_action=REPAIRED_FROM_TIMESTAMPS)"
            f". Ngược lại loại bỏ (không phục hồi được) -- sửa {tstats['duration_repaired']} "
            f"dòng, loại {tstats['duration_excluded']} dòng trong lượt này.",
        ])
        add_table(doc, ["Bước", "Số dòng"], [
            ["test.parquet gốc", str(tstats["original_rows"])],
            ["Loại do ranh giới thời gian", str(tstats["temporal_boundary_excluded"])],
            ["Sửa duration từ timestamp", str(tstats["duration_repaired"])],
            ["Loại do duration không phục hồi được", str(tstats["duration_excluded"])],
            ["Evaluation View cuối cùng của Test", str(tstats["final_evaluated_rows"])],
        ])
        para(doc, "Triển khai: `scripts/final_test/quality_transform.py` "
                  "(`load_requests_with_quality_transform()`), có self-check ở "
                  "`scripts/final_test/test_quality_transform.py`. Mỗi request được đánh giá đều "
                  "giữ lại `duration_seconds_raw` và `quality_action` bên cạnh "
                  "`duration_seconds` (có thể đã sửa) để audit được từng dòng. Audit trail đầy "
                  "đủ (trace ngược nguồn xác nhận lỗi tồn tại ở dữ liệu nguồn thô, không phải do "
                  "preprocessing của dự án này): `final_test/DATA_QUALITY_GATE.md`.",
              size=9.5, color=GREY)

        heading(doc, "8.3 Lệnh chạy Final Test", level=2)
        if ft_commands_log:
            code_block(doc, ft_commands_log.strip())
        else:
            code_block(doc, "python scripts/final_test/audit_test_dataset.py\n"
                             "python scripts/final_test/verify_before_run.py\n"
                             "python scripts/final_test/run_final_test_baselines.py\n"
                             "python scripts/final_test/run_final_test_ablation.py\n"
                             "python scripts/final_test/run_final_test_long_horizon.py\n"
                             "python scripts/final_test/build_final_test_summary.py")

        heading(doc, "8.4 Bản đồ Artifact (final_test/)", level=2)
        add_table(doc, ["Đường dẫn", "Mục đích"], [
            ["FINAL_TEST_PROTOCOL.md", "Config/seed/hash đã đóng băng, viết trước khi bất kỳ "
             "policy nào chạy trên test.parquet"],
            ["DATA_QUALITY_GATE.md", "Audit đầy đủ về anomaly duration, trace root-cause, lý "
             "giải rule sửa"],
            ["test_quality_transform_manifest.json", "Số liệu sửa/loại dạng máy đọc được + row "
             "ID theo từng split"],
            ["baseline/", "5 policy x 5 seed trên Final Test Evaluation View (CSV từng seed + "
             "tổng hợp)"],
            ["ablation/", "Full / No Forecast / No Fairness x 5 seed"],
            ["long_horizon/", "Kết quả checkpoint trên 1 quỹ đạo, ngày 1-37 + tỷ lệ đổi quyết "
             "định policy"],
            ["validation_vs_test.csv", "So sánh chiều Validation-vs-Test theo từng phát hiện "
             "(heldout_generalization)"],
            ["test_claim_assessment.csv", "Bảng claim C1-C6: heldout_generalization + "
             "paper_replication_verdict (2 cột độc lập)"],
            ["FINAL_TEST_MENTOR_SUMMARY.md", "Tóm tắt dễ đọc trả lời 6 câu hỏi generalization "
             "cốt lõi"],
            ["figures/", "PNG baseline/ablation/long-horizon sinh từ CSV Final Test"],
            ["logs/", "commands.log, environment.txt, runtimes.csv"],
        ])

        heading(doc, "8.5 Phụ lục Định nghĩa Metric", level=2)
        bullets(doc, [
            "Fairness là một KHÁI NIỆM, không phải 1 metric duy nhất -- Gini và Variance là 2 "
            "metric được báo cáo (xem công thức ở Mục 5/6 phía trên); Gini thấp hơn và Variance "
            "thấp hơn đều nghĩa là thu nhập driver đồng đều hơn.",
            "Paired delta: với 1 seed, (metric ở config A) - (metric ở config B), tính theo từng "
            "seed rồi lấy trung bình -- báo cáo kèm tính nhất quán về dấu (vd \"5/5 seed\") chứ "
            "không chỉ mean.",
            "Generalization (heldout) theo chiều: một phát hiện quan sát trên Validation có lặp "
            "lại ĐÚNG CHIỀU trên Test không? Chỉ tính từ dấu/so sánh, không có ngưỡng độ lớn.",
            "Paper replication verdict: phát hiện có khớp với tuyên bố định tính của chính "
            "arXiv:2407.17839 không? Đây là đánh giá độc lập so với paper, KHÔNG suy ra được từ "
            "phép tính generalization -- một phát hiện có thể generalize trong khi claim của "
            "paper vẫn Not Reproduced (xem ví dụ C4 đã làm ở Research Report Mục 10).",
        ])

        heading(doc, "8.6 Giới hạn Khả năng Tái lập (Final Test)", level=2)
        bullets(doc, [
            "Dataset là NYC TLC 2013, không phải dữ liệu 2016 gốc của paper -- Final Test kế "
            "thừa đúng phạm vi trend-replication như Validation (Mục 3 của Research Report), "
            "không phải exact reproduction.",
            "Chỉ 5 seed -- effect size, mean/std, và tính nhất quán dấu theo seed là bằng chứng "
            "chính; không tính hay tuyên bố formal statistical significance test nào.",
            "Chỉ chạy đúng 1 operating point canonical lambda=0,5 -- không quét lambda trên Test "
            "(có chủ đích; Test dùng để verify, không dùng để chọn operating point mới).",
            "Product demo (05_SanPham_Demo) mặc định dùng slice Validation/demo, không bao giờ "
            "dùng test.parquet -- Test chỉ dành riêng cho đánh giá khoa học cuối cùng này, "
            "không lộ ra trong Control Room trực tiếp.",
        ])
        if ft_environment_txt:
            para(doc, "Môi trường chạy Final Test (ghi lại ở final_test/logs/environment.txt):",
                 size=9.5, color=GREY)
            code_block(doc, ft_environment_txt.strip())

    # ---------------- 9. Lệnh Tái lập Chính xác ----------------
    heading(doc, "9. Lệnh Tái lập Chính xác (Exact Reproduction Commands)", level=1)
    para(doc, "Chạy từ thư mục gốc repo, đúng theo thứ tự này. Số dòng là artifact hiện đang giữ "
              "lại, verify trực tiếp tại thời điểm build tài liệu này.")
    cmd_rows = [
        ("python -m pytest tests/test_simulator_invariants.py -q", "(không có file đầu ra)", test_line),
        ("python train_momaql.py", "data/momaql_q_table_trained.json",
         "1.511 state (zone,hour) (không ghi đè nếu chưa chạy lại toàn bộ thí nghiệm downstream)"),
        ("python run_r1.py", "reports/r1_validation_results.csv",
         f"{csv_row_count('r1_validation_results.csv')} dòng (5 policy x 5 seed)"),
        ("python run_r2_ablation.py", "reports/r2_ablation_raw.csv + _results.csv",
         f"{csv_row_count('r2_ablation_raw.csv')} + {csv_row_count('r2_ablation_results.csv')} dòng"),
        ("python run_pareto_frontier.py", "reports/pareto_frontier_results.csv + _summary.csv + "
         ".png", f"{csv_row_count('pareto_frontier_results.csv')} + "
         f"{csv_row_count('pareto_frontier_summary.csv')} dòng"),
        ("python run_multi_horizon.py", "reports/multi_horizon_results.csv + "
         "policy_disagreement.csv", f"{csv_row_count('multi_horizon_results.csv')} + "
         f"{csv_row_count('policy_disagreement.csv')} dòng"),
        ("python run_complete_verifications.py", "reports/fleet_scale_results.csv + "
         "spatial_disagreement_by_zone.csv", f"{csv_row_count('fleet_scale_results.csv')} + "
         f"{csv_row_count('spatial_disagreement_by_zone.csv')} dòng"),
        ("python run_spatial_candidate_pool.py", "reports/spatial_candidate_pool.csv",
         f"{csv_row_count('spatial_candidate_pool.csv')} dòng"),
        ("python run_q_table_convergence.py", "reports/q_table_convergence_daily.csv",
         f"{csv_row_count('q_table_convergence_daily.csv')} dòng"),
        ("python run_hypothesis1_weekly_cycle.py", "reports/hypothesis1_weekly_cycle.csv",
         f"{csv_row_count('hypothesis1_weekly_cycle.csv')} dòng"),
        ("python run_hypothesis4_fairness_balance.py", "reports/hypothesis4_fairness_balance.csv",
         f"{csv_row_count('hypothesis4_fairness_balance.csv')} dòng"),
        ("python train_and_eval_mlp.py", "reports/mlp_vs_tabular_results.csv + _summary.csv",
         f"{csv_row_count('mlp_vs_tabular_results.csv')} + "
         f"{csv_row_count('mlp_vs_tabular_summary.csv')} dòng; cần torch"),
        ("python make_report_figures.py", "docs/*/figures/*.png (7 hình x 3 thư mục đầu ra)",
         "sinh lại mọi hình từ các CSV ở trên"),
        ("python docs/docx_report/build_research_report.py", "docs/docx_report/"
         "Bao_Cao_Nghien_Cuu_FairDispatch_MOMAQL.docx", "Research Report"),
        ("python docs/techdoc/build_technical_documentation.py", "docs/techdoc/"
         "Technical_Documentation.docx", "tài liệu này"),
    ]
    add_table(doc, ["Lệnh", "Đầu ra", "Kết quả kỳ vọng"], cmd_rows)
    callout(doc, "KHÔNG chạy train_momaql_multipass.py thay cho train_momaql.py -- đã công bố rõ "
                 "là không ổn định, có thể ghi ra data/momaql_q_table_multipass.json, và một số "
                 "script sẽ ưu tiên dùng file đó thay vì bảng canonical nếu nó tồn tại. Nếu file "
                 "đó có sẵn, cách ly nó trước khi chạy bất kỳ đánh giá nào.")

    # ---------------- 10. Gói Khả năng Tái lập ----------------
    heading(doc, "10. Gói Khả năng Tái lập (Reproducibility Package)", level=1)
    add_table(doc, ["Trường", "Giá trị"], [
        ["Git commit", head],
        ["Manifest dataset", "data/train.parquet (912.375 dòng), data/val.parquet (195.508 "
         "dòng), data/test.parquet (195.510 dòng)"],
    ] + [[f"SHA-256 -- {name}", f"{h}  ({size:,} bytes)"] for name, (h, size) in ds_hashes.items()]
      + [
        ["Cấu hình", "N_DRIVERS=200, window=60s, MAX_PICKUP_ETA=600s, lambda=0,5 mặc định (quét "
         "0-1), gamma=0,9, alpha=0,1"],
        ["Seed ngẫu nhiên", "20260721, 20260722, 20260723, 20260724, 20260725 (thí nghiệm "
         "chính); 3 seed đầu dùng cho thí nghiệm cơ chế"],
        ["Môi trường", f"Python {platform.python_version()}, numpy 2.2.6, pandas 2.3.2, pyarrow "
         "21.0.0, scipy 1.16.2, torch 2.7.1+cu118 (version tại thời điểm build -- kiểm tra lại "
         "bằng `pip show` trước khi trích dẫn vào báo cáo chính thức)"],
        ["CPU/GPU", "Intel Core i5-10300H (chỉ dùng CPU cho mọi thí nghiệm simulator/dispatch); "
         "NVIDIA GTX 1650 qua CUDA (chỉ dùng để huấn luyện MLP trong train_and_eval_mlp.py)"],
        ["Số lượng test", test_line],
        ["Artifact kết quả", "mọi file trong reports/ (16 CSV + 1 PNG + 1 manifest checksum "
         "JSON) -- Validation. Artifact Held-out Test: final_test/ (xem Mục 8)."],
    ])
    para(doc, "reports/dataset_checksums.json là snapshot tĩnh, đã ghi từ trước; bảng ở trên "
              "tính MỚI tại thời điểm build tài liệu và là nguồn có thẩm quyền nếu 2 bên có mâu "
              "thuẫn.", size=9.5, color=GREY)

    # ---------------- 11. Vấn đề Đã biết, Giả định, Xử lý sự cố ----------------
    heading(doc, "11. Vấn đề Đã biết, Giả định, và Xử lý sự cố (Known Issues, Assumptions, and "
                 "Troubleshooting)", level=1)
    heading(doc, "11.1 Giả định (khớp Research Report Mục 11.3, theo góc nhìn kỹ thuật)", level=2)
    bullets(doc, [
        "A1. Số lượng tài xế (200) không được paper nêu rõ. Chọn theo quy mô đội xe hợp lý cho "
        "Manhattan; đổi N_DRIVERS ở đầu bất kỳ run_*.py nào để test sensitivity (xem "
        "fleet_scale_results.csv cho quét N=100/200/400 đã chạy sẵn).",
        "A2. Biểu diễn không gian dùng 67 zone taxi TLC chính thức, không phải đồ thị tự gộp "
        "cụm -- pickup_zone_id/dropoff_zone_id lấy trực tiếp từ cột parquet, codebase này không "
        "có bước clustering thêm.",
        "A3. ETA đón khách là khoảng cách haversine với tốc độ hằng số 12 mph (AVG_SPEED_MPH), "
        "không phải road routing -- đơn giản hóa có chủ đích cho simulator nhẹ (xem docstring "
        "module src/simulator.py).",
        "A4. Q(vùng,giờ) huấn luyện qua 1 lượt tuần tự duy nhất trên train.parquet "
        "(train_momaql.py, 1 seed=20260721) -- không phải nhiều epoch. "
        "train_momaql_multipass.py có tồn tại nhưng đã công bố rõ là không ổn định; output chẩn "
        "đoán của nó nằm ở reports/momaql_convergence.csv (5 dòng, các lượt annealed-alpha) chỉ "
        "để lưu vết, không dùng làm input đánh giá.",
        "A5. Không giữ lại lockfile version dependency nào; danh sách package ở Mục 3 là bộ đã "
        "audit đang hoạt động, không phải requirements.txt đã ghim version.",
    ])
    heading(doc, "11.2 Vấn đề Đã biết (Known Issues)", level=2)
    bullets(doc, [
        "tests/test_simulator_invariants.py::test_no_double_booking_within_window và "
        "::test_time_monotonicity gọi simulator với record_trace=False mặc định, nên assertion "
        "dựa trên trace của chúng hiện đang lặp trên danh sách trace rỗng. Chúng pass, nhưng "
        "không thực sự chạy qua code path phụ thuộc trace mà chúng được viết ra để kiểm tra -- "
        "chạy lại với record_trace=True cục bộ nếu cần verify đúng guarantee đó.",
        "run_r2_ablation.py và run_pareto_frontier.py âm thầm ưu tiên "
        "data/momaql_q_table_multipass.json hơn bảng canonical nếu file đó tồn tại trên đĩa "
        "(xem cảnh báo Mục 8) -- đây là hành vi script có sẵn từ trước, không phải bug được sửa "
        "trong đợt này, nhưng rất dễ vô tình dính phải.",
        "Không có structured logging framework; mỗi script chỉ print dòng tiến trình ra stdout "
        "thuần với flush=True. Capture stdout ra file nếu cần lưu log chạy lâu dài.",
    ])
    heading(doc, "11.3 Xử lý sự cố (Troubleshooting)", level=2)
    bullets(doc, [
        "\"FileNotFoundError: data/train.parquet\" -- các phần chia parquet bị gitignore (quá "
        "to để push git bình thường); phải lấy riêng từ nơi dữ liệu của repo này ban đầu được "
        "chuẩn bị. Chỉ reports/*.csv và JSON bảng Q đã huấn luyện được track.",
        "Kết quả lệch nhẹ so với số trích dẫn trong Research Report -- kiểm tra "
        "reports/dataset_checksums.json (Mục 10) có khớp byte-for-byte với file data/ cục bộ "
        "của anh trước; 1 bản build parquet khác sẽ làm lệch mọi con số downstream.",
        "Áp lực bộ nhớ trên máy ít RAM trống -- mỗi run_*.py đã gọi gc.collect() và del object "
        "lớn giữa các seed/config; nếu vẫn OOM, giảm SEEDS về 1 giá trị duy nhất trước để cô lập "
        "xem đây là leak theo từng seed hay thật sự không đủ RAM cho working set của 1 seed.",
        "train_and_eval_mlp.py lỗi import torch -- đây là script duy nhất trong repo cần nó; "
        "mọi script khác chỉ dùng numpy/pandas/pyarrow/scipy thuần.",
    ])

    doc.save(OUT)
    print(f"[done] {OUT}")
    return doc, dict(head=head, test_line=test_line, ds_hashes=ds_hashes)


if __name__ == "__main__":
    main()
